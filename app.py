"""
app.py — Indexed Content Pipeline — Streamlit UI
Indexed brand theme: dark navy, cyan accent, clean cards.
All clients: Smart Fog, Upvolt, OE Partners, RecNation, InspectMind, Veriheal.
"""
from __future__ import annotations
import os, re, sys, time, subprocess
from datetime import datetime
from pathlib import Path

import streamlit as st
from docx import Document

ROOT = Path(__file__).parent

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Content Pipeline — Indexed",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Indexed brand CSS ─────────────────────────────────────────────────────────
st.markdown("""<style>
:root {
  --bg:    #0b1112;
  --bg2:   #0f1c1c;
  --bg3:   #132525;
  --bd:    #1e3535;
  --ac:    #1ab8b0;
  --ac-hi: #22d4cb;
  --acd:   rgba(26,184,176,.14);
  --t:     #ffffff;
  --m:     #8faaa8;
  --s:     #4a6866;
  --r:     12px;
  --pill:  100px;
}

/* ── Base ── */
.stApp { background: var(--bg) !important; }
html, body, [class*="css"] {
  font-family: system-ui, -apple-system, "Segoe UI", sans-serif !important;
}
#MainMenu, footer, header,
[data-testid="stToolbar"],
[data-testid="stDecoration"],
[data-testid="stStatusWidget"] { display: none !important; }

/* ── Sidebar ── */
[data-testid="stSidebar"] {
  background: #080f0f !important;
  border-right: 1px solid var(--bd) !important;
}
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span { color: var(--m) !important; }
[data-testid="stSidebar"] strong { color: var(--ac) !important; }

/* ── Typography ── */
h1, h2, h3 {
  color: var(--t) !important;
  font-weight: 800 !important;
  letter-spacing: -0.02em !important;
}
p, li { color: var(--m) !important; line-height: 1.7 !important; }
label, [data-testid="stWidgetLabel"] p {
  color: var(--s) !important;
  font-size: .75rem !important;
  font-weight: 700 !important;
  letter-spacing: .08em !important;
  text-transform: uppercase !important;
}
[data-testid="stCaption"] p { color: var(--s) !important; font-size: .75rem !important; }
code { background: var(--bg2) !important; color: var(--ac) !important;
       border-radius: 4px !important; padding: .1em .4em !important; }

/* ── Buttons ── */
.stButton > button {
  border-radius: var(--pill) !important;
  font-weight: 700 !important;
  letter-spacing: .01em !important;
  transition: all .15s !important;
  padding: .45rem 1.4rem !important;
}
.stButton > button[kind="primary"] {
  background: var(--ac) !important;
  color: #0b1112 !important;
  border: none !important;
}
.stButton > button[kind="primary"]:hover { filter: brightness(1.12) !important; }
.stButton > button[kind="primary"]:disabled {
  background: var(--bd) !important;
  color: var(--s) !important;
  opacity: .6 !important;
}
.stButton > button[kind="secondary"] {
  background: transparent !important;
  color: var(--m) !important;
  border: 1px solid var(--bd) !important;
}
.stButton > button[kind="secondary"]:hover {
  border-color: var(--ac) !important;
  color: var(--t) !important;
}

/* ── Download button ── */
.stDownloadButton > button {
  border-radius: var(--pill) !important;
  background: var(--acd) !important;
  color: var(--ac) !important;
  border: 1px solid var(--ac) !important;
  font-weight: 600 !important;
}
.stDownloadButton > button:hover {
  background: var(--ac) !important;
  color: #0b1112 !important;
}

/* ── Text input ── */
[data-testid="stTextInput"] input {
  background: var(--bg2) !important;
  border: 1px solid var(--bd) !important;
  border-radius: var(--r) !important;
  color: var(--t) !important;
}
[data-testid="stTextInput"] input:focus {
  border-color: var(--ac) !important;
  box-shadow: 0 0 0 2px var(--acd) !important;
  outline: none !important;
}
[data-testid="stTextInput"] input::placeholder { color: var(--s) !important; }

/* ── Selectbox ── */
[data-baseweb="select"] > div {
  background: var(--bg2) !important;
  border: 1px solid var(--bd) !important;
  border-radius: var(--r) !important;
}
[data-baseweb="select"] span,
[data-baseweb="select"] div { color: var(--t) !important; }
[data-baseweb="popover"] {
  background: var(--bg2) !important;
  border: 1px solid var(--bd) !important;
  border-radius: var(--r) !important;
  box-shadow: 0 8px 32px rgba(0,0,0,.6) !important;
}
[data-baseweb="menu"] li { color: var(--m) !important; }
[data-baseweb="menu"] li:hover { background: var(--bg3) !important; color: var(--t) !important; }

/* ── Tabs ── */
[data-baseweb="tab-list"] {
  background: transparent !important;
  border-bottom: 1px solid var(--bd) !important;
}
[data-baseweb="tab"] {
  color: var(--m) !important;
  font-weight: 600 !important;
  background: transparent !important;
}
[data-baseweb="tab"][aria-selected="true"] {
  color: var(--ac) !important;
  border-bottom: 2px solid var(--ac) !important;
}
[data-testid="stTabContent"] { padding-top: 1.5rem !important; }

/* ── Progress bar ── */
[data-testid="stProgress"] > div {
  background: var(--bd) !important;
  height: 3px !important;
  border-radius: var(--pill) !important;
}
[data-testid="stProgress"] > div > div {
  background: var(--ac) !important;
  border-radius: var(--pill) !important;
}

/* ── Status / expander / code ── */
[data-testid="stStatus"],
[data-testid="stExpander"] {
  background: var(--bg2) !important;
  border: 1px solid var(--bd) !important;
  border-radius: var(--r) !important;
}
[data-testid="stCode"] pre {
  background: #060d0d !important;
  border: 1px solid var(--bd) !important;
  border-radius: var(--r) !important;
  color: var(--m) !important;
  font-size: .75rem !important;
}

/* ── Dividers / alerts ── */
hr, [data-testid="stDivider"] { border-color: var(--bd) !important; }
[data-baseweb="notification"] { border-radius: var(--r) !important; border: none !important; }
</style>""", unsafe_allow_html=True)

# ── Client registry ───────────────────────────────────────────────────────────
CLIENTS: dict[str, tuple[str, str, str, str]] = {
    # key: (label, icon, description, flow)
    "smart-fog":   ("Smart Fog",   "💧", "Industrial Humidity Control", "plan"),
    "upvolt":      ("Upvolt",      "☀️", "Solar Energy Installations",  "plan"),
    "oe-partners": ("OE Partners", "🎯", "Lean Six Sigma Training",     "plan"),
    "recnation":   ("RecNation",   "🚐", "RV & Boat Storage",           "plan"),
    "inspectmind": ("InspectMind", "🏗️", "AEC Plan Review",             "plan"),
    "veriheal":    ("Veriheal",    "🌿", "Medical Cannabis Platform",   "url"),
}

PIPELINE_STAGES = [
    "Scrape / Parse content",
    "GSC keyword check",
    "Ahrefs keyword research",
    "Competitor analysis",
    "NeuronWriter NLP",
    "Brief generation",
    "Article generation",
]

# ── Session state ─────────────────────────────────────────────────────────────
_DEFAULTS: dict = {
    "client":       "smart-fog",
    "article_path": None,
    "brief_path":   None,
    "is_running":   False,
    "last_name":    None,
    "last_time":    None,
    "error":        None,
    "stage_num":    0,
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ── Helpers ───────────────────────────────────────────────────────────────────

def _extract_para_text(para) -> str:
    try:
        from lxml import etree
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        parts = []
        for child in para._p:
            local = etree.QName(child).localname
            if local == "hyperlink":
                r_id = child.get(f"{{{R}}}id")
                anchor = "".join(t.text or "" for t in child.findall(f".//{{{W}}}t"))
                if r_id and hasattr(para.part, "rels") and r_id in para.part.rels:
                    url = para.part.rels[r_id].target_ref
                    parts.append(f"[{anchor}]({url})" if url else anchor)
                else:
                    parts.append(anchor)
            elif local == "r":
                parts.append("".join(t.text or "" for t in child.findall(f".//{{{W}}}t")))
        return "".join(parts).strip()
    except Exception:
        return (para.text or "").strip()


def _read_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        lines = []
        for p in doc.paragraphs:
            txt = _extract_para_text(p)
            if not txt:
                continue
            s = p.style.name
            if   "Heading 1" in s: lines.append(f"# {txt}")
            elif "Heading 2" in s: lines.append(f"## {txt}")
            elif "Heading 3" in s: lines.append(f"### {txt}")
            elif "List"      in s: lines.append(f"- {txt}")
            else:                  lines.append(txt)
        return "\n\n".join(lines)
    except Exception as e:
        return f"_Could not render document: {e}_"


def _word_count(path: Path) -> int:
    try:
        doc = Document(str(path))
        return len(" ".join(p.text for p in doc.paragraphs).split())
    except Exception:
        return 0


def _latest_output(client: str, kind: str) -> Path | None:
    d = ROOT / "outputs" / client
    if not d.exists():
        return None
    files = [f for f in d.glob("*.docx") if not f.name.startswith("_")]
    files = ([f for f in files if f.name.lower().startswith("brief")] if kind == "brief"
             else [f for f in files if not f.name.lower().startswith("brief")])
    return sorted(files, key=lambda f: f.stat().st_mtime, reverse=True)[0] if files else None


def _load_rows(client: str) -> list[tuple[int, str, str]]:
    try:
        import config as cfg
        profile = cfg.load_client_profile(client)
        cp = profile.get("content_plan", {})
        if not cp:
            return []
        xls = os.path.join(cfg.ROOT_DIR, cp.get("excel_path", ""))
        if not xls or not os.path.exists(xls):
            return []
        from pipeline.content_plan_parser import parse_content_plan
        rows = parse_content_plan(xls, row_index=None, tab_name=None, profile=profile)
        return [
            (int(r.get("content_plan_row", i)),
             r.get("content_plan_topic", "")[:80],
             r.get("content_plan_primary_keyword", ""))
            for i, r in enumerate(rows)
            if r.get("content_plan_primary_keyword")
        ]
    except Exception as e:
        st.error(f"Could not load content plan: {e}")
        return []


def _run_pipeline(args: list[str]):
    env = {**os.environ, "PYTHONUNBUFFERED": "1"}
    proc = subprocess.Popen(
        [sys.executable, "-u", str(ROOT / "main.py")] + args,
        stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
        text=True, bufsize=1, cwd=str(ROOT), env=env,
    )
    for line in proc.stdout:
        yield line.rstrip()
    proc.wait()


def _dl_btn(path: Path, label: str):
    with open(path, "rb") as f:
        st.download_button(
            label=label, data=f.read(), file_name=path.name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

# ── Sidebar — stage tracker ───────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""<div style="padding:.5rem 0 1rem;">
      <span style="font-size:1.1rem;font-weight:900;letter-spacing:-.02em;color:#fff;">INDEXED</span>
    </div>""", unsafe_allow_html=True)
    st.caption("PIPELINE STAGES")
    st.markdown('<div style="height:.25rem"></div>', unsafe_allow_html=True)

    cur = st.session_state.stage_num
    for i, stage in enumerate(PIPELINE_STAGES, 1):
        if st.session_state.is_running and i == cur:
            st.markdown(f"**→ {stage}**")
        elif cur > i:
            st.caption(f"✓ {stage}")
        else:
            st.caption(f"· {stage}")

    if st.session_state.last_name:
        st.divider()
        st.caption("LAST GENERATED")
        st.caption(f"📄 {st.session_state.last_name}")
        if st.session_state.last_time:
            st.caption(f"🕐 {st.session_state.last_time}")

# ── Header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div style="display:flex;align-items:center;gap:.85rem;
     padding:1.2rem 0 1.1rem;border-bottom:1px solid #1e3535;margin-bottom:2rem;">
  <span style="font-size:1.4rem;font-weight:900;letter-spacing:-.03em;
               color:#fff;font-family:system-ui;">INDEXED</span>
  <span style="background:rgba(26,184,176,.1);color:#1ab8b0;font-size:.65rem;
               font-weight:700;letter-spacing:.1em;text-transform:uppercase;
               padding:.2rem .75rem;border:1px solid rgba(26,184,176,.3);
               border-radius:100px;">Content Pipeline</span>
  <span style="flex:1;"></span>
  <span style="color:#4a6866;font-size:.78rem;">Automated brief &amp; article generation</span>
</div>
""", unsafe_allow_html=True)

# ── Client selector ───────────────────────────────────────────────────────────
st.markdown(
    '<p style="color:#4a6866;font-size:.72rem;font-weight:700;'
    'letter-spacing:.1em;text-transform:uppercase;margin-bottom:.6rem;">Client</p>',
    unsafe_allow_html=True,
)

client_cols = st.columns(len(CLIENTS))
for col, (key, (label, icon, desc, flow)) in zip(client_cols, CLIENTS.items()):
    with col:
        is_sel = (st.session_state.client == key)
        bd  = "#1ab8b0"               if is_sel else "#1e3535"
        bg  = "rgba(26,184,176,.07)"  if is_sel else "#0f1c1c"
        tc  = "#ffffff"               if is_sel else "#8faaa8"
        dc  = "#1ab8b0"               if is_sel else "#4a6866"
        st.markdown(f"""
        <div style="background:{bg};border:1px solid {bd};border-radius:12px;
                    padding:.85rem .75rem;margin-bottom:.4rem;min-height:76px;">
          <div style="font-size:1.2rem;margin-bottom:.2rem;">{icon}</div>
          <div style="font-size:.82rem;font-weight:700;color:{tc};
                      margin-bottom:.1rem;">{label}</div>
          <div style="font-size:.67rem;color:{dc};">{desc}</div>
        </div>""", unsafe_allow_html=True)
        if st.button(
            "✓ Selected" if is_sel else "Select",
            key=f"cl_{key}",
            use_container_width=True,
            type="primary" if is_sel else "secondary",
            disabled=st.session_state.is_running,
        ):
            st.session_state.client = key
            st.rerun()

st.markdown('<div style="height:1.2rem"></div>', unsafe_allow_html=True)
st.divider()
st.markdown('<div style="height:.4rem"></div>', unsafe_allow_html=True)

# ── Input section ─────────────────────────────────────────────────────────────
client_key = st.session_state.client
label, icon, desc, flow = CLIENTS[client_key]
run_args: list[str] = []

if flow == "plan":
    st.markdown(
        '<p style="color:#4a6866;font-size:.72rem;font-weight:700;'
        'letter-spacing:.1em;text-transform:uppercase;margin-bottom:.3rem;">Article</p>',
        unsafe_allow_html=True,
    )
    rows = _load_rows(client_key)
    if rows:
        options = {(topic or f"Row {idx}"): idx for idx, topic, _ in rows}
        sel_label = st.selectbox(
            "article",
            list(options.keys()),
            label_visibility="collapsed",
            disabled=st.session_state.is_running,
        )
        sel_row = options[sel_label]
        sel_kw  = next((kw for idx, _, kw in rows if idx == sel_row), "")
        st.caption(f"Row {sel_row}  ·  {sel_kw}")
        run_args = ["--client", client_key, "--row", str(sel_row)]
    else:
        row_num = st.number_input(
            "Row number", min_value=0, value=0, step=1,
            disabled=st.session_state.is_running,
        )
        run_args = ["--client", client_key, "--row", str(int(row_num))]

else:
    st.markdown(
        '<p style="color:#4a6866;font-size:.72rem;font-weight:700;'
        'letter-spacing:.1em;text-transform:uppercase;margin-bottom:.3rem;">Article URL</p>',
        unsafe_allow_html=True,
    )
    url = st.text_input(
        "url", placeholder="https://www.veriheal.com/blog/...",
        label_visibility="collapsed",
        disabled=st.session_state.is_running,
    )
    _u = url.strip()
    if _u:
        if "veriheal.com" not in _u:
            st.warning("URL should be a veriheal.com article.")
        else:
            run_args = ["--client", "veriheal", "--url", _u]
    else:
        st.caption("Enter a Veriheal article URL to continue.")

# ── Generate button ───────────────────────────────────────────────────────────
st.markdown('<div style="height:.6rem"></div>', unsafe_allow_html=True)
col_btn, col_hint = st.columns([2, 5])
with col_btn:
    generate = st.button(
        "⚡  Generate Brief + Article",
        type="primary",
        disabled=(not run_args or st.session_state.is_running),
        use_container_width=True,
    )
with col_hint:
    if st.session_state.is_running:
        st.caption("⚙️  Pipeline running — please wait…")
    else:
        st.caption("⏱  Avg generation time: 6–8 min")

# ── Pipeline execution ────────────────────────────────────────────────────────
if generate and run_args and not st.session_state.is_running:
    st.session_state.is_running   = True
    st.session_state.article_path = None
    st.session_state.brief_path   = None
    st.session_state.error        = None
    st.session_state.stage_num    = 0

    start = time.time()
    os.makedirs(ROOT / "outputs" / client_key, exist_ok=True)

    log_lines: list[str] = []
    err_lines: list[str] = []
    pbar     = st.progress(0.0, text="Starting…")
    log_ph   = st.empty()

    with st.status("Pipeline running…", expanded=True) as status:
        for line in _run_pipeline(run_args):
            if not line.strip():
                continue

            m = re.search(r'Stage\s+(\d+)\s*/\s*(\d+)\s*[—\-]+\s*(.*)', line)
            if m:
                n, tot, name = int(m.group(1)), int(m.group(2)), m.group(3).strip()
                st.session_state.stage_num = n
                pbar.progress((n - 1) / tot, text=f"Stage {n}/{tot} — {name}")
                continue

            if "===" in line:
                continue

            low = line.lower()
            if any(k in low for k in ("error", "failed", "critical", "traceback", "exception")):
                err_lines.append(line.strip())

            log_lines.append(line.strip())
            log_ph.code("\n".join(log_lines[-18:]), language=None)

        elapsed = time.time() - start
        pbar.progress(1.0, text="Complete")

        art_found = _latest_output(client_key, "article")
        if err_lines and not art_found:
            st.session_state.error = err_lines[-1]
            status.update(label=f"❌ Failed after {elapsed:.0f}s", state="error", expanded=True)
        else:
            status.update(label=f"✅ Done in {elapsed:.0f}s", state="complete", expanded=False)

    _art = _latest_output(client_key, "article")
    _brf = _latest_output(client_key, "brief")
    st.session_state.article_path = str(_art) if _art else None
    st.session_state.brief_path   = str(_brf) if _brf else None
    st.session_state.is_running   = False
    st.session_state.stage_num    = 0

    if _art:
        st.session_state.last_name = _art.stem[:50]
        st.session_state.last_time = datetime.now().strftime("%d %b %Y, %H:%M")

    st.rerun()

# ── Error ─────────────────────────────────────────────────────────────────────
if st.session_state.error and not st.session_state.article_path:
    st.error(f"Pipeline failed: {st.session_state.error}")
    st.caption("Check the log above for the full trace, then re-run.")

# ── Outputs ───────────────────────────────────────────────────────────────────
if st.session_state.article_path or st.session_state.brief_path:
    st.divider()

    art_path = Path(st.session_state.article_path) if st.session_state.article_path else None
    brf_path = Path(st.session_state.brief_path)   if st.session_state.brief_path   else None

    wc       = _word_count(art_path) if art_path and art_path.exists() else 0
    wc_label = f" · {wc:,} words" if wc else ""

    tab_a, tab_b = st.tabs([f"📄 Article{wc_label}", "📋 Brief"])

    with tab_a:
        if art_path and art_path.exists():
            c1, c2 = st.columns([1, 4])
            with c1:
                _dl_btn(art_path, "⬇ Download Article")
            with c2:
                st.caption(f"`{art_path.name}`")
                if st.session_state.last_time:
                    st.caption(f"Generated {st.session_state.last_time}")
            st.divider()
            st.markdown(_read_docx(art_path))
        else:
            st.warning("No article file found.")

    with tab_b:
        if brf_path and brf_path.exists():
            c1, c2 = st.columns([1, 4])
            with c1:
                _dl_btn(brf_path, "⬇ Download Brief")
            with c2:
                st.caption(f"`{brf_path.name}`")
            st.divider()
            st.markdown(_read_docx(brf_path))
        else:
            st.warning("No brief file found.")
