"""
article_generator.py — Generates the Phase 2 optimised article as a .docx file.
Includes add_hyperlink() for embedding live links directly in body copy.
"""

from __future__ import annotations
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic
import config
from pipeline.context_loader import load_all_context, format_context_for_prompt


def add_hyperlink(paragraph, url: str, anchor_text: str):
    """
    Embed a live clickable hyperlink into a python-docx paragraph.
    If url is empty (UNCONFIRMED), writes anchor_text as italic plain text
    with no URL suffix — the QA report flags these for editor resolution.
    """
    if not url:
        run = paragraph.add_run(anchor_text)
        run.italic = True
        return None

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "0563C1")
    rPr.append(color_el)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = anchor_text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _load_system_prompt(profile: dict) -> str:
    sp_path = profile.get("_system_prompt_path", "")
    if not sp_path or not os.path.exists(sp_path):
        raise FileNotFoundError(f"System prompt not found: {sp_path}")
    with open(sp_path, "r", encoding="utf-8") as f:
        return f.read()


def _load_article_rules(profile: dict) -> str:
    """
    Load client-specific article writing rules from clients/{client}/article_rules.md.
    These rules are injected into the article generation user message and replace the
    former hardcoded Veriheal voice/style/linking/disclaimer block.
    Falls back to an empty string if not found — the system_prompt.md should cover the gap.
    """
    rules_path = os.path.join(
        profile.get("_client_dir", ""),
        "article_rules.md",
    )
    if os.path.exists(rules_path):
        with open(rules_path, "r", encoding="utf-8") as f:
            return f.read()
    return ""


def _read_brief(brief_path: str) -> str:
    """
    Read the brief .docx and return its text content.
    Includes paragraph text AND table rows so internal link tables are captured.
    """
    try:
        from docx.oxml.ns import qn as _qn
        doc = Document(brief_path)
        parts = []
        # Iterate body children in document order so tables appear in context
        body = doc.element.body
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                text = child.text or ""
                # Collect all run text
                runs_text = "".join(
                    r.text for r in child.iter()
                    if r.tag.endswith("}t") and r.text
                )
                if runs_text.strip():
                    parts.append(runs_text)
            elif tag == "tbl":
                for row in child.iter():
                    if row.tag.endswith("}tr"):
                        cells = []
                        for cell in row:
                            if cell.tag.endswith("}tc"):
                                cell_text = "".join(
                                    t.text for t in cell.iter()
                                    if t.tag.endswith("}t") and t.text
                                )
                                cells.append(cell_text.strip())
                        row_text = " | ".join(cells)
                        if row_text.strip():
                            parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        raise RuntimeError(f"Could not read brief file: {e}")


def _build_user_message(
    url: str,
    article_data: dict,
    brief_text: str,
    keyword_data: dict,
    nlp_terms: dict,
    context_text: str,
    profile: dict | None = None,
) -> str:

    primary_kw = keyword_data.get("primary_keyword", "N/A") or "N/A"

    _target_wc = None
    if not nlp_terms.get("enabled") or not nlp_terms.get("terms"):
        nlp_block = "NeuronWriter integration pending. No NLP terms available."
    else:
        terms = nlp_terms["terms"]
        _target_wc = nlp_terms.get("target_word_count")
        target_wc_display = _target_wc or "N/A"
        terms_formatted = "\n".join([
            f"  - {t['term']}: {t['usage_range']} (used by {t['usage_pc']}% of top pages)"
            for t in terms
        ])
        nlp_block = f"NeuronWriter target word count: {target_wc_display}\nNLP terms to use naturally (do not stuff):\n{terms_formatted}"

    # Final safety net: if _target_wc still not set (NeuronWriter pending or API didn't return it),
    # extract from brief text. The brief generator always writes "Target word count: NNN" in the
    # NeuronWriter Guidance section, so this works regardless of API availability.
    if not _target_wc:
        _brief_wc_match = re.search(r"Target word count[:\s]+(\d+)", brief_text, re.IGNORECASE)
        if _brief_wc_match:
            _target_wc = int(_brief_wc_match.group(1))
            print(f"  Word count ceiling: {_target_wc} words (extracted from brief)")

    # Build word count ceiling instruction
    if _target_wc:
        _wc_ceiling = int(_target_wc * 1.1)
        _wc_instruction = (
            f"\n--- WORD COUNT CEILING — NON-NEGOTIABLE ---\n\n"
            f"The NeuronWriter target for this article is {_target_wc} words. "
            f"The article MUST NOT exceed {_wc_ceiling} words (10% tolerance). "
            f"This is a hard ceiling, not a suggestion.\n\n"
            f"Before writing each section, check its target word count in the brief "
            f"(listed as 'Target words:' under each H2). Do not exceed it.\n\n"
            f"When you must cut to meet the ceiling, cut in this order:\n"
            f"1. Mechanism explanation (how the technology works) — cut to 2-3 sentences\n"
            f"2. Background context and definitions\n"
            f"3. Never cut: comparison data, trade-off analysis, industry application specifics, "
            f"specifications with citations\n\n"
        )
    else:
        _wc_instruction = ""

    # Extract internal links from brief text for explicit injection
    internal_links_block = "No internal links specified in brief."
    if "Internal Links to Add" in brief_text:
        try:
            links_section = brief_text.split("Internal Links to Add")[1]
            # Cut off at next section
            for end_marker in ["Disclaimer", "NeuronWriter", "---"]:
                if end_marker in links_section:
                    links_section = links_section.split(end_marker)[0]
                    break
            internal_links_block = links_section.strip()
        except Exception:
            pass

    _profile = profile or {}
    _domain = _profile.get("domain", "")
    _client_name = _profile.get("client_name", "Client")
    _article_rules = _load_article_rules(_profile)
    _is_net_new = article_data.get("is_net_new", False)

    # Content plan context block — injected for net-new (content-plan) flows
    if _is_net_new:
        _funnel = article_data.get("content_plan_funnel", "not specified")
        _cp_type = article_data.get("content_plan_type", "not specified")
        _cluster = article_data.get("content_plan_cluster", "not specified")
        _priority = article_data.get("content_plan_priority", "not specified")
        _landing_page = article_data.get("content_plan_landing_page", "not specified")
        _strategy_notes = article_data.get("content_plan_strategy_notes", "")
        _ai_visibility_notes = article_data.get("content_plan_ai_visibility_notes", "")
        _ai_block = (
            f"\nAI Visibility Notes:\n{_ai_visibility_notes}\n"
            if _ai_visibility_notes else ""
        )
        article_content_plan_block = (
            f"\n--- CONTENT PLAN CONTEXT ---\n\n"
            f"Funnel Stage: {_funnel}\n"
            f"Content Type: {_cp_type}\n"
            f"Cluster: {_cluster}\n"
            f"Priority: {_priority}\n"
            f"Associated Landing Page: {_landing_page}\n\n"
            f"Strategy Notes:\n{_strategy_notes if _strategy_notes else '(not provided)'}\n"
            + _ai_block
        )
    else:
        article_content_plan_block = ""

    # Build disclaimer block from profile
    _disclaimers = _profile.get("disclaimers", [])
    if _disclaimers:
        _disclaimer_block = "\n\n".join(_disclaimers)
        _disclaimer_instruction = (
            f"--- DISCLAIMERS — COPY VERBATIM ---\n\n"
            f"The following disclaimer(s) must appear at the end of the article body, "
            f"after the CTA and before the FAQ. Copy them word-for-word — do not paraphrase, "
            f"truncate, or omit any.\n\n"
            + "\n\n".join(f"Disclaimer {i+1} (copy exactly):\n{d}" for i, d in enumerate(_disclaimers))
        )
    else:
        _disclaimer_instruction = "(No disclaimers required for this client.)"

    # Domain-specific URL prepend instruction
    _url_prepend = (
        f"If a URL below is a partial path (e.g. /blog/example), prepend https://www.{_domain}."
        if _domain else
        "If a URL below is a partial path, prepend the client domain."
    )

    return f"""You are now beginning Phase 2 — article generation for {_client_name}.

--- PHASE 1 BRIEF ---

{brief_text}

--- ORIGINAL ARTICLE DATA ---

URL: {url}
Title: {article_data.get('title')}
H1: {article_data.get('h1')}
Word count: {article_data.get('word_count')}
Primary keyword: {primary_kw}
Current headings: {json.dumps([h['text'] for h in article_data.get('headings', [])], indent=2)}
Body text (first 3000 chars):
{article_data.get('body_text', '')[:3000]}
{article_content_plan_block}
--- NEURONWRITER NLP TERMS ---

{nlp_block}

--- PROJECT CONTEXT ---

{context_text}

--- INTERNAL LINKS — MANDATORY ---

You MUST embed EVERY internal link listed below in the article body — not
a selection, not a subset, ALL of them. If the list contains 4 links, all
4 must appear in the article. If it contains 5, all 5 must appear.
Before finishing the article, count the links in the list below and verify
each one appears in the body. Missing any link is a QA failure.
Use the EXACT anchor text and EXACT URL shown below — do not modify or invent URLs.
Format every link precisely as [HYPERLINK: anchor text | url].
Do not skip any link. Do not substitute your own URLs.
{_url_prepend}

{internal_links_block}

{_article_rules}

{_disclaimer_instruction}

{_wc_instruction}--- STRUCTURE ---

You must only write sections explicitly listed in the brief under
'Sections to Add' and 'Sections to Change', plus the existing
sections listed under 'H2 Recommendations'. Do not invent, add,
or expand sections beyond what the brief specifies. If you think
a section would improve the article, do not add it — the brief
is the complete specification. Adding unrequested sections is a
compliance failure.

H2 AND H3 PLACEMENT LABELS: Every H2 and H3 must carry a placement label derived from the brief's structure. Rules:
- Read the brief's H2/H3 order to determine where each section sits relative to existing live headings
- The label sits on its own line, immediately after the heading line, before any body text or H3s
- It must never appear inside a paragraph, mid-sentence, or after body text has started
- Format in italics using exactly one of these forms:
  *(New H2 — insert after intro)*
  *(New H2 — insert between "[H2 before]" and "[H2 after]")*
  *(New H2 — insert before [section name])*
  *(Replaces existing H2: "[original heading]")*
  *(Existing H2 — updated content replaces current section in full)*
  *(New H3 — insert under "[parent H2]", between "[H3 before]" and "[H3 after]")*
  *(New H3 — insert under "[parent H2]", before first existing H3)*
- The label text must reflect the brief exactly — use the heading names from the brief, not invented descriptions
- For H3s, always name the parent H2 and the surrounding H3s from the brief so the editor knows exactly where to place it
- Missing or misplaced labels are a compliance failure
- Labels are for the WordPress editor and are removed after placement

PARAGRAPH FLOW: Every H2 and H3 section must open with at least one full body paragraph before any table, bullet list, or numbered list appears. This paragraph should introduce the section topic and connect it to what came before. Never place a table or list as the first element directly after a heading.

NO PREVIEW REDUNDANCY: Opening paragraphs must NOT restate or preview the data that bullet clusters or lists immediately below will deliver. The opening paragraph's job is to establish WHY this dimension matters for the reader's decision — not to summarise the conclusions that the structured content already shows. If you find yourself writing "Steam systems require more maintenance while dry fog extends service intervals..." before a bullet cluster that says exactly that, delete the prose and keep only the bullet cluster. Every claim must appear once. Writing it in prose and then again in a bullet is redundancy, not depth.

COMPARISON FORMATTING — BULLET CLUSTERS: When the brief calls for a comparison of 3 or more items across multiple attributes (technology types, product categories, equipment classes, or any multi-row comparison a table would show), format each comparison dimension as a bullet cluster. A bullet cluster places all compared items under one named dimension heading. Format each bullet as: **Item Name:** one sentence. The bold item name and colon are mandatory on every bullet — no exceptions, no item may be written without **bold**: format. Use a colon after the bold name — not a hyphen, not an em dash. Rules: (1) every compared item must appear in every cluster — no gaps; (2) one sentence per bullet — two sentences makes it a paragraph, not a scannable comparison; (3) the comparison dimensions must be established at the start of the section and held consistently — do not introduce new dimensions mid-section or vary which items get covered. The brief will specify the comparison criteria — use exactly those as cluster dimensions, do not substitute. For 2-item comparisons, use structured paragraphs instead of bullet clusters.

--- INSTRUCTION ---

Do NOT write a metadata table. The pipeline generates it automatically from
the brief — if you write one it will be duplicated. Begin your response
directly with the H1.

Execute the brief exactly. Write every section as specified.
Do not output any commentary, voice calibration notes, compliance checks,
or chat window blocks. Output the article and nothing else.
Begin with H1, then body sections in order, then disclaimers (if required), then FAQ.
CTA PLACEMENT: The CTA (call to action) must appear immediately after the
final body section and before the disclaimers and FAQ. Never place the CTA after the FAQ.

CTA HYPERLINK FORMAT: Every CTA that invites the reader to contact, request a quote, or speak with an engineer MUST use the [HYPERLINK: anchor text | url] format on the anchor phrase. The anchor text alone becomes the clickable link — the surrounding sentence provides context. Example of correct format:
  Speak with a Smart Fog engineer about humidification requirements for your facility by visiting [HYPERLINK: speak with a Smart Fog engineer | https://www.smartfog.com/contact-us/].
Writing the CTA as plain prose with no HYPERLINK marker is a formatting failure — the link will not be embedded in the document.
"""


def _fill_formatted_paragraph(para, text: str):
    """Fill an existing paragraph with bold-aware runs (splits on **)."""
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (idx % 2 == 1)


_BULLET_ABSTRACT_NUM_ID = 100
_DECIMAL_ABSTRACT_NUM_ID = 101


def _ensure_numbering_part(doc):
    """Bootstrap a numbering.xml part if the document has none. Returns the part."""
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        numbering_xml = (
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b'</w:numbering>'
        )
        part = Part(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_xml,
            doc.part.package,
        )
        doc.part.relate_to(
            part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
        )
        numbering_part = doc.part.numbering_part
    return numbering_part


def _apply_bullet_numbering(doc, para) -> None:
    """
    Register a bullet abstractNum (id=100) the first time it is needed,
    then attach numId 100 to the paragraph via <w:numPr>.
    """
    from lxml import etree
    numbering_part = _ensure_numbering_part(doc)
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = numbering_part._element

    if not root.findall(f"w:abstractNum[@w:abstractNumId='{_BULLET_ABSTRACT_NUM_ID}']", nsmap):
        root.append(etree.fromstring(
            f'<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:abstractNumId="{_BULLET_ABSTRACT_NUM_ID}">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0">'
            '<w:start w:val="1"/><w:numFmt w:val="bullet"/>'
            '<w:lvlText w:val="&#x2022;"/><w:lvlJc w:val="left"/>'
            '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
            '<w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr>'
            '</w:lvl></w:abstractNum>'
        ))
        root.append(etree.fromstring(
            f'<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:numId="{_BULLET_ABSTRACT_NUM_ID}">'
            f'<w:abstractNumId w:val="{_BULLET_ABSTRACT_NUM_ID}"/></w:num>'
        ))

    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(_BULLET_ABSTRACT_NUM_ID))
    numPr.append(ilvl)
    numPr.append(numId_el)
    pPr.append(numPr)


def _new_numbered_list_id(doc) -> int:
    """
    Register the decimal abstractNum (id=101) if not yet present, then create
    a brand-new <w:num> instance referencing it and return its numId.
    Each call produces a fresh counter that starts at 1 — use one per list.
    """
    from lxml import etree
    numbering_part = _ensure_numbering_part(doc)
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = numbering_part._element

    if not root.findall(f"w:abstractNum[@w:abstractNumId='{_DECIMAL_ABSTRACT_NUM_ID}']", nsmap):
        root.append(etree.fromstring(
            f'<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:abstractNumId="{_DECIMAL_ABSTRACT_NUM_ID}">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0">'
            '<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
            '<w:lvlText w:val="%1."/><w:lvlJc w:val="left"/>'
            '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
            '</w:lvl></w:abstractNum>'
        ))

    existing_nums = root.findall("w:num", nsmap)
    new_id = max((int(n.get(qn("w:numId"), 0)) for n in existing_nums), default=0) + 1
    root.append(etree.fromstring(
        f'<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:numId="{new_id}">'
        f'<w:abstractNumId w:val="{_DECIMAL_ABSTRACT_NUM_ID}"/></w:num>'
    ))
    return new_id


def _apply_numbered_list(para, num_id: int) -> None:
    """Attach a specific numId to a paragraph for numbered-list rendering."""
    pPr = para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId_el)
    pPr.append(numPr)


def _add_formatted_paragraph(doc, text: str, style: str = "Normal"):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(6)
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (idx % 2 == 1)
    return para


def _run_qa_check(
    article_text: str,
    brief_text: str,
    nlp_terms: dict,
    keyword_data: dict,
    confirmed_external_count: int,
    bare_domain_links: list,
    client: anthropic.Anthropic,
    profile: dict | None = None,
    article_rules: str = "",
) -> dict:
    """
    Runs a QA check on the generated article against the brief.
    Returns a dict with critical_failures (list) and report (str).
    """
    # Deterministic pre-check: flag zero confirmed external links before Claude runs.
    pre_critical_failures: list[str] = []
    pre_qa_prefix = ""
    if confirmed_external_count == 0:
        _msg = (
            "EXTERNAL LINKS — ZERO CONFIRMED: Article contains no confirmed external "
            "links. At least 2-3 external citations to authoritative sources are "
            "required. All external links are either missing or marked UNCONFIRMED."
        )
        pre_critical_failures.append(_msg)
        pre_qa_prefix = f"CRITICAL FAILURES:\n- {_msg}\n\n"

    # Deterministic pre-check: flag bare-domain external links as non-critical.
    if bare_domain_links:
        bare_domain_flags = []
        for anchor, url in bare_domain_links:
            bare_domain_flags.append(
                f"[HYPERLINK: {anchor} | {url}] — bare domain, no article path"
            )
        pre_qa_prefix += (
            "\nNON-CRITICAL FLAGS (Python pre-check):\n"
            + "\n".join(f"- {f}" for f in bare_domain_flags)
            + "\n"
        )

    terms = nlp_terms.get("terms", [])
    terms_summary = "\n".join([
        f"  - {t['term']}: {t['usage_range']}"
        for t in terms[:20]
    ]) if terms else "No NLP terms available."

    _profile = profile or {}
    _client_name = _profile.get("client_name", "Client")
    blocked_domains = list(_profile.get("blocked_external_sources", [
        "leafwell.com", "nuggmd.com", "leafly.com",
        "greenhealthdocs.com", "docmj.com", "quickmedcards.com",
        "weedmaps.com",
    ]))
    acceptable_sources = list(_profile.get("allowed_external_sources", [
        "pubmed.ncbi.nlm.nih.gov", "nih.gov", "cdc.gov", "fda.gov",
        "nida.nih.gov", "mayoclinic.org", "hopkinsmedicine.org",
        "health.harvard.edu", "jamanetwork.com", "nejm.org",
        "thelancet.com", "pmc.ncbi.nlm.nih.gov",
    ]))

    # Use client-specific article rules when available; fall back to hardcoded Veriheal standards
    if article_rules:
        _voice_standards_block = f"--- {_client_name.upper()} VOICE AND STYLE STANDARDS ---\n\n{article_rules}"
    else:
        _voice_standards_block = f"""--- {_client_name.upper()} VOICE AND STYLE STANDARDS ---

Voice:
- Conversational but credible. Plain language. Second person ("you") throughout.
- Short paragraphs: 2-3 sentences ideal, maximum 5 sentences.
- Cause-and-effect connections explained simply, not academically.
- Helpful and informative. Not robotic or mechanical.
- Down-to-earth tone. Never stiff or corporate.

Style rules:
- No em dashes (—). Flag every occurrence.
- Oxford comma required in lists of three or more.
- "Cannabis" preferred over "weed" or "marijuana" in body copy.
- No slang: no "stoner", "pot", "black market", "getting high" in medical context.
- Spell out abbreviations on first use: "tetrahydrocannabinol (THC)".
- Paragraphs should be 2-5 sentences. Flag any single-sentence paragraphs used excessively.
- No double spaces after periods.
- Present perfect tense preferred where applicable.
- No AI-sounding phrases: "it's worth noting", "it's important to understand",
  "in conclusion", "in summary", "delve", "comprehensive", "crucial", "essential",
  "game-changer", "furthermore", "moreover", "showcasing", "let's explore".

Links:
- At least one internal link (to a {_client_name} URL) must be present.
- At least one external link to an acceptable source must be present.
- No links to blocked competitor domains.
- External links must go to appropriate authoritative sources only.
- Flag any [HYPERLINK: anchor | UNCONFIRMED] markers as needing resolution."""

    qa_prompt = f"""You are a senior editor at {_client_name} conducting a QA audit of an optimised article.
Evaluate the article against the brief and {_client_name}'s editorial standards. Return a structured QA report.

--- BRIEF ---
{brief_text}

--- ARTICLE ---
{article_text}

--- NLP TERMS ---
{terms_summary}

--- BLOCKED EXTERNAL DOMAINS ---
{chr(10).join(blocked_domains)}

--- ACCEPTABLE EXTERNAL SOURCES ---
Only these domains are acceptable for external links. All others are a critical failure:
{chr(10).join(acceptable_sources)}
Or equivalent authoritative institutions (.gov, .edu, peer-reviewed journals).

{_voice_standards_block}

--- QA INSTRUCTIONS ---

Check the following and report on each:

CRITICAL CHECKS (failure blocks article saving):
1. DISCLAIMER: Are all required disclaimers present in the article (if any are required for this client)?
2. INTERNAL LINKS: Are all internal links from the brief present with correct anchor text and URLs?
3. EXTERNAL LINKS - BLOCKED DOMAINS: Do any external links go to blocked competitor domains?
4. EXTERNAL LINKS - SOURCE QUALITY: Do any external links go to random blogs, non-authoritative sources, or inappropriate sources?
5. METADATA TABLE: Is the metadata table present with revised title, meta description, slug, and H1?
6. EXTERNAL LINK MANDATORY: Search the article text for any [HYPERLINK:
anchor | url] where the url does NOT contain "{_profile.get('domain', 'client-domain.com')}". This includes
urls that say "UNCONFIRMED". If you find even ONE such marker — confirmed
url or UNCONFIRMED — this check PASSES. Only mark this as a CRITICAL
FAILURE if there is literally no [HYPERLINK:] marker at all with a
non-client url (including UNCONFIRMED). Do not fail this check
because a link is UNCONFIRMED — UNCONFIRMED counts as present.
7. MINIMUM INTERNAL LINK: Does the article have at least one confirmed
internal client link embedded as a hyperlink?

NON-CRITICAL CHECKS (flagged but does not block saving):
8. UNCONFIRMED LINKS: Only flag a link as UNCONFIRMED if the url
field literally contains the word "UNCONFIRMED" (e.g. [HYPERLINK:
some anchor | UNCONFIRMED]). List any such markers that need a real
URL resolved before publishing. Bare-domain links are handled
separately by a pre-check and do not need to be flagged here.
9. WORD COUNT: Estimate the article word count. Is it within the brief's recommended range?
10. SECTIONS TO CHANGE: Were all sections in the brief's Sections to Change addressed?
11. SECTIONS TO ADD: Were all new sections in the brief present and placed correctly?
12. EM DASHES: Flag any em dashes (—) found in the article body paragraphs. Do NOT flag em dashes in headings derived from the content plan title.
13. AI PHRASES: Flag any AI-sounding phrases from the voice standards list above found in BODY PARAGRAPHS ONLY. Do not flag words that appear in H1, H2, H3 headings, or the metadata table — headings come from the content plan and are not generated prose. Only flag body paragraph occurrences.
14. PARAGRAPH LENGTH: Flag any sections with excessive single-sentence paragraphs.
15. ABBREVIATIONS: Flag any technical abbreviations that were not spelled out on first use in body copy (relevant to the client's industry — ignore cannabis-specific abbreviations for non-cannabis clients).
16. VOICE QUALITY: Does the article read as conversational, patient-focused, and plain-language? Flag any sections that feel mechanical, academic, or robotic.
17. NLP TERMS: Are the NLP terms used naturally without stuffing? Flag any obvious forced insertions.
18. SLANG: Flag any use of "pot", "stoner", or "black market". Note that "weed"
and "marijuana" are acceptable in titles, H1s, H2s, and occasionally in body
copy for natural variation — do not flag these.

OUTPUT FORMAT — use exactly this structure:

CRITICAL FAILURES:
- [list each critical failure, or write "None" if all passed]

NON-CRITICAL FLAGS:
- [list each non-critical issue, or write "None" if all passed]

QA VERDICT: [PASS or FAIL]

QA NOTES FOR EDITOR:
[3-4 sentences summarising overall quality, what reads well, and what the editor should fix before publishing]
"""

    response = client.messages.create(
        model=config.ANTHROPIC_MODEL,
        max_tokens=2000,
        messages=[{"role": "user", "content": qa_prompt}],
    )

    qa_text = ""
    for block in response.content:
        if hasattr(block, "text") and block.text:
            qa_text += block.text

    # Prepend any deterministic pre-failures so they appear in the QA report.
    qa_text = pre_qa_prefix + qa_text

    # Parse critical failures from Claude's response.
    critical_failures = list(pre_critical_failures)  # start with pre-checked failures
    in_critical = False
    for line in qa_text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("CRITICAL FAILURES:"):
            in_critical = True
            continue
        if stripped.startswith("NON-CRITICAL FLAGS:"):
            in_critical = False
            continue
        if in_critical and stripped.startswith("-"):
            failure = stripped[1:].strip()
            if not failure.lower().startswith("none") and failure not in critical_failures:
                critical_failures.append(failure)

    verdict = "FAIL" if critical_failures else "PASS"

    return {
        "critical_failures": critical_failures,
        "verdict": verdict,
        "report": qa_text,
    }


def _append_qa_report(doc: Document, qa_report: str) -> None:
    """Append the QA report to the end of the article document."""
    doc.add_page_break()
    heading = doc.add_heading("QA REPORT — EDITOR REVIEW", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "This section is for editorial review only. Remove before publishing."
    ).italic = True
    doc.add_paragraph("")
    for line in qa_report.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped in ("CRITICAL FAILURES:", "NON-CRITICAL FLAGS:",
                        "QA VERDICT:", "QA NOTES FOR EDITOR:"):
            para = doc.add_paragraph(style="Normal")
            run = para.add_run(stripped)
            run.bold = True
        elif stripped.startswith("-"):
            doc.add_paragraph(stripped[1:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)


def _fill_table_cell(cell, text: str) -> None:
    """
    Write text into a docx table cell with bold-aware and hyperlink-aware rendering.

    Handles two marker types, interleaved correctly:
      - [HYPERLINK: anchor | url]  → live clickable link via add_hyperlink()
      - **text**                   → bold run

    Plain text segments between markers are written as normal runs.
    Never writes raw asterisks or raw [HYPERLINK:] syntax into the cell.
    """
    import re as _re
    para = cell.paragraphs[0]
    # Clear any default content python-docx adds
    for run in para.runs:
        run.text = ""

    # Split on [HYPERLINK: anchor | url] markers, keeping the groups
    _hl_pat = _re.compile(r'\[HYPERLINK:\s*([^\|]+?)\s*\|\s*([^\]]+?)\s*\]')
    segments = _hl_pat.split(text)
    # split() with a capturing group produces:
    # [plain, anchor1, url1, plain, anchor2, url2, ...]

    idx = 0
    while idx < len(segments):
        if idx % 3 == 0:
            # Plain text segment — render with ** bold handling
            plain = segments[idx]
            if plain:
                bold_parts = plain.split("**")
                for bidx, bpart in enumerate(bold_parts):
                    if not bpart:
                        continue
                    run = para.add_run(bpart)
                    run.bold = (bidx % 2 == 1)
        elif idx % 3 == 1:
            # Anchor text — paired with url at idx+1
            anchor = segments[idx].strip()
            url = segments[idx + 1].strip() if (idx + 1) < len(segments) else ""
            if url.upper() == "UNCONFIRMED":
                url = ""
            add_hyperlink(para, url, anchor)
            idx += 1  # skip the url segment; outer loop advances past it
        idx += 1


def _parse_and_write_doc(article_text: str, url: str, output_path: str, qa_report: str = "", client_name: str = "Content") -> None:
    """
    Parse Claude's article output and write it to a .docx file.
    Converts [HYPERLINK: anchor | url] markers into live embedded links.
    """
    # Strip Claude's internal commentary blocks
    lines_raw = article_text.split("\n")
    cleaned_lines = []
    skip_mode = False
    for line in lines_raw:
        stripped = line.strip()
        if any(phrase in stripped.lower() for phrase in [
            "chat window updates",
            "chat window log",
            "chat window flags",
            "chat window",
        ]):
            skip_mode = True
        if skip_mode:
            continue
        if any(phrase in stripped.lower() for phrase in [
            "voice calibration",
            "primary keyword placement",
            "phase 2 proceeding",
            "output 2 delivered",
            "i'll now execute phase",
            "brief compliance check",
            "executing phase 2",
        ]):
            continue
        cleaned_lines.append(line)
    article_text = "\n".join(cleaned_lines)

    doc = Document()

    # Metadata header
    title_para = doc.add_heading(f"{client_name.upper()} OPTIMISED ARTICLE", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"URL: {url}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    doc.add_paragraph("")

    lines = article_text.split("\n")
    i = 0
    _current_num_id: int | None = None   # numId for the active numbered list
    _prev_was_numbered: bool = False      # was the previous non-empty line a numbered item?
    _prev_was_heading: bool = False      # was the previous non-empty line a heading?
    _prev_was_bullet: bool = False       # was the previous non-empty line a bullet item?
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Reset numbered-list state for any non-numbered content line
        _is_numbered_item = (
            len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)"
        )
        if not _is_numbered_item:
            _prev_was_numbered = False

        if stripped == "---" or stripped.startswith("|---|"):
            i += 1
            continue

        if stripped.startswith("### "):
            _prev_was_bullet = False
            _prev_was_heading = True
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            _prev_was_bullet = False
            _prev_was_heading = True
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            _prev_was_bullet = False
            _prev_was_heading = True
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            if _prev_was_bullet:
                doc.add_paragraph("")  # blank line after bullet block
            _prev_was_bullet = False
            _prev_was_heading = False
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not all(c in "-| " for c in row):
                    table_lines.append(row)
                i += 1
            if table_lines:
                headers = [c.strip() for c in table_lines[0].strip("|").split("|")]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for j, h in enumerate(headers):
                    hdr_cells[j].text = h
                    for run in hdr_cells[j].paragraphs[0].runs:
                        run.bold = True
                for row_line in table_lines[1:]:
                    cells = [c.strip() for c in row_line.strip("|").split("|")]
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(cells):
                        if j < len(row_cells):
                            _fill_table_cell(row_cells[j], cell_text)
                doc.add_paragraph("")
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            if not _prev_was_bullet and not _prev_was_heading:
                doc.add_paragraph("")  # blank line before bullet block (body text → bullets only)
            _prev_was_heading = False
            _prev_was_bullet = True
            para = doc.add_paragraph(style="List Bullet")
            _apply_bullet_numbering(doc, para)
            _fill_formatted_paragraph(para, stripped[2:])
            i += 1
            continue

        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)":
            if _prev_was_bullet:
                doc.add_paragraph("")  # blank line after bullet block
            _prev_was_bullet = False
            _prev_was_heading = False
            # New list boundary: allocate a fresh numId so counting restarts at 1
            if not _prev_was_numbered:
                _current_num_id = _new_numbered_list_id(doc)
            text_without_prefix = stripped[2:].lstrip()
            para = doc.add_paragraph(style="Normal")
            _fill_formatted_paragraph(para, text_without_prefix)
            _apply_numbered_list(para, _current_num_id)
            _prev_was_numbered = True
            i += 1
            continue

        # Normalise markdown links [text](url) → [HYPERLINK: text | url]
        if "](" in stripped and "[" in stripped:
            import re
            stripped = re.sub(
                r'\[([^\]]+)\]\((https?://[^\)]+)\)',
                lambda m: f"[HYPERLINK: {m.group(1)} | {m.group(2)}]",
                stripped,
            )

        if "[HYPERLINK:" in stripped:
            if _prev_was_bullet:
                doc.add_paragraph("")  # blank line after bullet block
            _prev_was_bullet = False
            _prev_was_heading = False
            para = doc.add_paragraph()
            parts = stripped.split("[HYPERLINK:")
            if parts[0]:
                para.add_run(parts[0])
            for part in parts[1:]:
                if "]" in part:
                    link_content, remainder = part.split("]", 1)
                    if "|" in link_content:
                        anchor, href = link_content.split("|", 1)
                        anchor = anchor.strip()
                        href = href.strip()
                        if href.upper() == "UNCONFIRMED":
                            add_hyperlink(para, "", anchor)
                        else:
                            add_hyperlink(para, href, anchor)
                    else:
                        para.add_run(f"[HYPERLINK:{link_content}]")
                    if remainder:
                        para.add_run(remainder)
                else:
                    para.add_run(f"[HYPERLINK:{part}")
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            if _prev_was_bullet:
                doc.add_paragraph("")  # blank line after bullet block
            _prev_was_bullet = False
            _prev_was_heading = False
            para = doc.add_paragraph(style="Normal")
            run = para.add_run(stripped.strip("**"))
            run.bold = True
            i += 1
            continue

        # Placement label line e.g. *(New H2 — insert after intro)*
        # Rendered as italic text; does not reset _prev_was_heading so that
        # a bullet block immediately after a heading+label still gets no blank paragraph.
        if stripped.startswith("*(") and stripped.endswith(")*"):
            para = doc.add_paragraph(style="Normal")
            run = para.add_run(stripped.strip("*"))
            run.italic = True
            i += 1
            continue

        # Plain body paragraph — close bullet block, then render normally
        if _prev_was_bullet:
            doc.add_paragraph("")  # blank line after bullet block
        _prev_was_bullet = False
        _prev_was_heading = False
        _add_formatted_paragraph(doc, stripped)
        i += 1

    if qa_report:
        _append_qa_report(doc, qa_report)

    doc.save(output_path)


_SERPER_ALLOWED_DOMAINS = {
    # PubMed / NIH / CDC / FDA
    "pubmed.ncbi.nlm.nih.gov", "pmc.ncbi.nlm.nih.gov", "ncbi.nlm.nih.gov",
    "nih.gov", "cdc.gov", "fda.gov", "nida.nih.gov",
    # Medical journals
    "jamanetwork.com", "nejm.org", "thelancet.com", "bmj.com",
    "annals.org", "acpjournals.org",
    # Health publishers
    "medicalnewstoday.com", "healthline.com", "webmd.com",
    "mayoclinic.org", "hopkinsmedicine.org", "health.harvard.edu",
    "clevelandclinic.org", "mountsinai.org",
    # Global health authorities
    "who.int", "emcdda.europa.eu",
    # News / wire services
    "reuters.com", "apnews.com", "bbc.com", "bbc.co.uk",
    "theguardian.com", "irishtimes.com", "independent.ie",
    # Irish / EU government and policy
    "gov.ie", "hse.ie", "hpra.ie", "cso.ie",
    "legislation.gov.uk", "eur-lex.europa.eu",
    # Drug policy / addiction / psychiatry
    "drugpolicy.org", "drugabuse.gov", "samhsa.gov", "psychiatry.org",
    "drugscience.org.uk",
    # Irish academic institutions
    "tcd.ie", "ucd.ie", "rcpi.ie",
    # Cannabis research and analytical sources
    "analyticalcannabis.com",       # lab-verified terpene/cannabinoid analytical content
    "projectcbd.org",               # non-profit CBD research and education organisation
    "cannabissciencetech.com",      # analytical/lab science content on cannabis chemistry
    "mjbizdaily.com",               # industry and market claims only — not for medical claims
    "frontiersin.org",              # open-access peer-reviewed journals, substantial cannabis output
    "mdpi.com",                     # open-access peer-reviewed publisher
    "researchgate.net",             # peer-reviewed study abstracts where PubMed has no direct page
}

_SERPER_BLOCKED_TERMS = {
    "shop", "store", "dispensary", "collective", "seeds",
    "weed", "420", "ganja",
}


def _is_serper_url_allowed(url: str) -> bool:
    """
    Returns True if the URL is from an authoritative source acceptable for
    citation. Blocks cannabis retail/blog domains; allows gov/edu/org and
    known health/news publishers.
    """
    import re as _re
    try:
        from urllib.parse import urlparse
        host = urlparse(url).netloc.lower().lstrip("www.")
    except Exception:
        return False

    # Always allow known authoritative domains (exact or subdomain match)
    for allowed in _SERPER_ALLOWED_DOMAINS:
        if host == allowed or host.endswith("." + allowed):
            return True

    # Always allow .gov and .edu TLDs
    if host.endswith(".gov") or host.endswith(".edu"):
        return True

    # Block domains containing cannabis/retail keywords
    host_parts = _re.split(r'[.\-]', host)
    if any(term in host_parts for term in _SERPER_BLOCKED_TERMS):
        return False

    # Allow remaining .org domains (WHO, advocacy, etc.)
    if host.endswith(".org"):
        return True

    # Everything else: block — keeps results to known quality sources
    return False


def _serper_search(query: str) -> list[str]:
    """
    Call Serper Google Search API and return up to 5 filtered organic result URLs.
    Filters out cannabis retail, dispensary, and blog domains before returning.
    Returns empty list on any error or if SERPER_API_KEY is not set.
    """
    if not config.SERPER_API_KEY:
        return []
    try:
        import requests as _requests
        resp = _requests.post(
            "https://google.serper.dev/search",
            headers={
                "X-API-KEY": config.SERPER_API_KEY,
                "Content-Type": "application/json",
            },
            json={"q": query, "num": 10},  # fetch more to allow for filtering
            timeout=10,
        )
        if not resp.ok:
            print(f"  WARNING: Serper search failed ({resp.status_code}) — falling back to Claude memory.")
            return []
        data = resp.json()
        all_urls = [r.get("link", "") for r in data.get("organic", []) if r.get("link")]
        filtered = [u for u in all_urls if _is_serper_url_allowed(u)]
        return filtered[:5]
    except Exception as e:
        print(f"  WARNING: Serper search error — {e}")
        return []


def _resolve_external_citations(
    article_text: str,
    client: anthropic.Anthropic,
    article_data: dict,
) -> str:
    """
    Pre-second-pass citation resolver.

    Scans for:
      (a) [HYPERLINK: anchor | UNCONFIRMED] markers
      (b) Unlinked sentences containing specific factual signals:
          percentages, named legislation, study/research attribution,
          named health/policy organisations, year-attributed findings.

    For each claim, calls Claude citation specialist (exact system + user
    prompts as specified). Confirmed URLs replace UNCONFIRMED markers or
    are embedded mid-sentence. Unresolved claims are left unchanged.

    Caps total claims at 8 (UNCONFIRMED markers take priority over
    unlinked sentences) to keep API call count bounded.
    """
    import re

    _CITE_SYSTEM = (
        "You are a citation specialist. Your only job is to return the single best "
        "confirmed URL for the claim provided. Return only the URL — no explanation, "
        "no markdown, no punctuation. If you are not confident the URL exists exactly "
        "as you recall it, return the single word UNCONFIRMED. Never fabricate or "
        "guess a URL.\n\n"
        "You have strong reliable knowledge of landmark studies in cannabis research. "
        "For well-documented findings like linalool's anxiolytic properties, CBD and "
        "pediatric epilepsy (Dravet syndrome), cannabis terpene interactions, and tobacco "
        "health effects, you should be able to provide confirmed PubMed or PMC URLs with "
        "high confidence. Err toward providing a confirmed URL when you have strong recall "
        "of a study rather than defaulting to UNCONFIRMED out of excessive caution. "
        "The URL you return must be for a source that is directly relevant to the article "
        "topic — do not return a URL that matches the anchor text but covers an unrelated subject."
    )

    _SOURCE_GUIDANCE = (
        "Choose the best authoritative source for the type of claim:\n\n"
        "For scientific/medical claims: prefer PubMed (https://pubmed.ncbi.nlm.nih.gov/) "
        "or PMC (https://pmc.ncbi.nlm.nih.gov/articles/) study pages. NIH, CDC, FDA, "
        "NIDA, Mayo Clinic, Johns Hopkins, Harvard Health, and JAMA are also acceptable.\n\n"
        "For legal/legislative claims: prefer official government websites (.gov.ie, "
        "legislation.gov.uk, eur-lex.europa.eu), official parliamentary records, or "
        "legal databases.\n\n"
        "For polling and statistics: the publishing organisation's official page — "
        "e.g. Irish Times/Ipsos poll pages, CSO Ireland (cso.ie), Eurobarometer.\n\n"
        "For policy and programme descriptions: official health authority pages "
        "(hpra.ie, hse.ie, gov.ie), WHO policy pages, or EMCDDA (emcdda.europa.eu).\n\n"
        "For news-based facts and documented events where no official source exists: "
        "authoritative news sources are acceptable — BBC, Irish Times, The Guardian, "
        "Reuters, Associated Press.\n\n"
        "Return only the URL or UNCONFIRMED."
    )

    _BLOCKED = {
        "leafly.com", "leafwell.com", "nuggmd.com", "weedmaps.com",
        "greenhealthdocs.com", "docmj.com", "quickmedcards.com",
    }

    _FACTUAL_PAT = re.compile(
        r'(?:'
        r'\d+\.?\d*\s*%|'                                            # percentages
        r'\d+\s+percent|'                                            # X percent
        r'\b(?:Act|Law|Bill|Regulation|Directive)\s+(?:of\s+)?\d{4}|'  # legislation
        r'\b(?:study|research|trial|survey|report)\s+'
        r'(?:found|showed?|suggests?|indicates?|published)|'         # study attribution
        r'\baccording\s+to\s+(?:the\s+)?[A-Z][A-Za-z]+|'           # "according to Org"
        r'\b(?:WHO|CDC|NIH|EMCDDA|HSE|FDA|NIDA|HRB)\b'             # named organisations
        r')',
        re.IGNORECASE,
    )

    claims: list[dict] = []

    # (a) UNCONFIRMED markers
    _unconf_pat = re.compile(r'\[HYPERLINK:([^\|]+)\|\s*UNCONFIRMED\]')
    for m in _unconf_pat.finditer(article_text):
        claims.append({
            "type": "unconfirmed",
            "anchor": m.group(1).strip(),
            "original": m.group(0),
        })

    # (b) Unlinked factual sentences — skip headings, table rows, hyperlink lines
    _seen: set[str] = set()
    for line in article_text.split("\n"):
        s = line.strip()
        if not s or "[HYPERLINK:" in s or s.startswith("|") or s.startswith("#"):
            continue
        if not _FACTUAL_PAT.search(s):
            continue
        # Extract first sentence from the paragraph line
        sent_m = re.match(r'[^.!?]+[.!?]', s)
        sentence = sent_m.group(0).strip() if sent_m else s[:200].strip()
        if sentence in _seen or "[HYPERLINK:" in sentence:
            continue
        _seen.add(sentence)
        claims.append({
            "type": "unlinked",
            "sentence": sentence,
            "full_line": s,
        })

    if not claims:
        print("  Citation lookup: no claims to check — skipping.")
        return article_text

    # Dynamic cap: raise to 12 for strain/terpene content categories
    # where claims are harder to resolve and volume is higher
    _STRAIN_SIGNALS = {
        "strain", "terpene", "terpenes", "flavonoid", "flavonoids",
        "cannabinoid", "cannabinoids", "indica", "sativa", "hybrid",
        "cultivar", "chemotype", "myrcene", "limonene", "linalool",
        "caryophyllene", "pinene", "ocimene", "terpinolene", "humulene",
    }
    _title_lower = article_data.get("title", "").lower()
    _h1_lower = article_data.get("h1", "").lower()
    _is_strain_article = any(
        s in _title_lower or s in _h1_lower
        for s in _STRAIN_SIGNALS
    )
    _claim_cap = 12 if _is_strain_article else 8

    # Prioritise UNCONFIRMED markers, apply dynamic cap
    unconf = [c for c in claims if c["type"] == "unconfirmed"]
    unlinked = [c for c in claims if c["type"] == "unlinked"][: max(0, _claim_cap - len(unconf))]
    claims = unconf + unlinked

    print(f"  Citation lookup: {len(claims)} claim(s) (cap: {_claim_cap}) — "
          f"{len(unconf)} UNCONFIRMED marker(s), {len(unlinked)} unlinked sentence(s).")

    _ANCHOR_SHARP_SYSTEM = (
        "You are a citation specialist. Your job is to rewrite a vague "
        "citation anchor into a specific, searchable form that will find "
        "a real study or authoritative source on PubMed, PMC, or a "
        "government health database. "
        "A good rewritten anchor names: the specific compound or "
        "intervention (e.g. CBD, myrcene, cannabidiol), the condition "
        "or outcome (e.g. acute pain, wound healing, inflammation), "
        "and ideally a year or study type (e.g. 2019 meta-analysis, "
        "PMC review). "
        "Return only the rewritten anchor — no explanation, no "
        "punctuation, no quotes. Keep it under 12 words. "
        "If the anchor already contains a named compound, condition, "
        "year, and study type, return it unchanged."
    )

    _VAGUE_PAT = re.compile(
        r'^(?:research\s+(?:suggests?|shows?|indicates?|found)|'
        r'studies?\s+(?:suggests?|shows?|indicates?|found)|'
        r'evidence\s+(?:suggests?|shows?|indicates?)|'
        r'(?:may|might|can)\s+(?:help|reduce|increase|improve|cause))',
        re.IGNORECASE,
    )

    def _needs_sharpening(anchor: str) -> bool:
        import re as _re
        # An anchor is specific enough if it contains at least one of:
        # - a year (4-digit number)
        # - a named compound, organism, or proper noun (capitalised
        #   mid-word or known cannabis/medical term)
        # - a named organisation or author
        # - a specific statistic (number + unit)
        _SPECIFIC = re.compile(
            r'\b\d{4}\b|'                          # year
            r'\b[A-Z][a-z]{2,}\b|'                 # proper noun / capitalised term
            r'\b(?:CBD|THC|CBG|CBN|THCA|CBDA|'
            r'linalool|myrcene|caryophyllene|'
            r'limonene|pinene|terpene|cannabinoid|'
            r'serotonin|dopamine|cortisol|GABA|'
            r'anandamide|endocannabinoid|'
            r'COX-2|cytokine|interleukin)\b|'      # named compound
            r'\b\d+\s*(?:mg|mg/kg|%|percent|'
            r'patients|participants|subjects)\b',  # statistic with unit
            re.IGNORECASE,
        )
        return not bool(_SPECIFIC.search(anchor))

    def _sharpen_anchor(anchor: str, context_sentence: str) -> str:
        """Call Claude to make a vague anchor specific. Returns original on error."""
        try:
            resp = client.messages.create(
                model=config.ANTHROPIC_MODEL,
                max_tokens=60,
                system=_ANCHOR_SHARP_SYSTEM,
                messages=[{"role": "user", "content": (
                    f"Anchor to sharpen: {anchor}\n"
                    f"Full sentence for context: {context_sentence}"
                )}],
            )
            sharpened = "".join(
                b.text for b in resp.content if hasattr(b, "text")
            ).strip().strip('"\'')
            return sharpened if sharpened else anchor
        except Exception:
            return anchor

    _article_topic = (
        f"{article_data.get('title', '').strip()} — {article_data.get('h1', '').strip()}"
    ).strip(" —")

    resolved = 0

    for claim in claims:
        if claim["type"] == "unconfirmed":
            raw_anchor = claim["anchor"]
            # Sharpen vague anchors before sending to citation specialist
            if _needs_sharpening(raw_anchor):
                sharp_anchor = _sharpen_anchor(raw_anchor, raw_anchor)
            else:
                sharp_anchor = raw_anchor
            claim_text = sharp_anchor
        else:
            claim_text = claim["sentence"]

        # Serper-augmented path: if key is set, search for candidate URLs
        # and append them to the prompt so Claude picks from real results
        # rather than relying on memory alone.
        serper_urls = _serper_search(claim_text) if config.SERPER_API_KEY else []
        if serper_urls:
            serper_block = (
                "Candidate URLs from search:\n"
                + "\n".join(f"  {u}" for u in serper_urls)
                + "\nSelect the most relevant URL from the list above that directly "
                "supports the claim. The article may cover multiple topics — a URL "
                "that supports the specific claim is acceptable even if it does not "
                "match the broader article topic. Return UNCONFIRMED only if none "
                "of the candidate URLs actually support the claim."
            )
            user_prompt = (
                f"Claim to cite: {claim_text}\n\n"
                + serper_block
                + "\n\n"
                + _SOURCE_GUIDANCE
            )
        else:
            # Memory-only path — identical to pre-Serper behaviour
            user_prompt = (
                f"Article topic: {_article_topic}\n"
                f"Find the best authoritative URL for this claim, relevant to the "
                f"article topic above: {claim_text}\n\n"
                + _SOURCE_GUIDANCE
            )

        try:
            resp = client.messages.create(
                model=config.ANTHROPIC_MODEL,
                max_tokens=120,
                system=_CITE_SYSTEM,
                messages=[{"role": "user", "content": user_prompt}],
            )
            raw = "".join(
                b.text for b in resp.content if hasattr(b, "text")
            ).strip().strip(".")

            # First non-empty line is the URL (or UNCONFIRMED)
            url = next((ln.strip() for ln in raw.splitlines() if ln.strip()), "")

            if url.upper() == "UNCONFIRMED" or not url.startswith("http"):
                continue
            if any(blocked in url for blocked in _BLOCKED):
                continue

            # Verify URL is live (returns 200) before accepting it.
            # Bare-domain URLs (no real path) are rejected immediately — they
            # signal the citation specialist defaulted to a homepage.
            try:
                from urllib.parse import urlparse as _up
                _parsed = _up(url)
                if _parsed.path.strip("/") == "":
                    print(f"  Citation lookup: bare-domain URL rejected — {url}")
                    continue
                import requests as _req
                _head = _req.head(url, timeout=6, allow_redirects=True,
                                  headers={"User-Agent": "Mozilla/5.0"})
                if _head.status_code not in (200, 301, 302, 303):
                    # Some servers block HEAD — try GET as fallback
                    _get = _req.get(url, timeout=6, allow_redirects=True,
                                    headers={"User-Agent": "Mozilla/5.0"}, stream=True)
                    if _get.status_code not in (200, 301, 302, 303):
                        print(f"  Citation lookup: URL returned {_get.status_code} — keeping UNCONFIRMED.")
                        continue
            except Exception as _ve:
                print(f"  Citation lookup: could not verify {url} ({_ve}) — keeping UNCONFIRMED.")
                continue

            resolved += 1

            if claim["type"] == "unconfirmed":
                article_text = article_text.replace(
                    claim["original"],
                    f"[HYPERLINK: {claim['anchor']} | {url}]",
                    1,
                )
            else:
                # Ask Claude to rewrite the sentence with a natural noun-phrase anchor.
                _embed_prompt = (
                    f"Rewrite this sentence to embed a hyperlink naturally. "
                    f"The URL to cite is: {url}\n\n"
                    f"Original sentence: {claim['sentence']}\n\n"
                    "Rules:\n"
                    "- The anchor text must be a noun phrase that describes what the linked source is about\n"
                    "- The anchor must sit on a noun or noun phrase, never on a verb phrase or attribution opener\n"
                    "- The full sentence must read correctly with the anchor text removed\n"
                    "- Do not add new factual claims — only restructure the existing sentence\n"
                    "- Use the format [HYPERLINK: anchor text | url] exactly\n"
                    "- Keep the sentence length similar to the original\n\n"
                    "Examples of correct embedding:\n"
                    "Original: 'Research suggests caffeine may influence how readily THC binds to cannabinoid receptors.'\n"
                    "Rewritten: 'Caffeine may influence [THC binding at cannabinoid receptors](url) according to current research.'\n"
                    "Formatted: 'Caffeine may influence [HYPERLINK: THC binding at cannabinoid receptors | url] according to current research.'\n\n"
                    "Return only the rewritten sentence with the embedded [HYPERLINK: anchor | url] marker."
                )
                try:
                    _embed_resp = client.messages.create(
                        model=config.ANTHROPIC_MODEL,
                        max_tokens=200,
                        messages=[{"role": "user", "content": _embed_prompt}],
                    )
                    _rewritten = "".join(
                        b.text for b in _embed_resp.content if hasattr(b, "text")
                    ).strip()
                    if _rewritten and "[HYPERLINK:" in _rewritten:
                        article_text = article_text.replace(
                            claim["full_line"],
                            claim["full_line"].replace(claim["sentence"], _rewritten, 1),
                            1,
                        )
                    else:
                        print(f"  Citation embed: rewrite returned no marker — skipping sentence.")
                except Exception as e:
                    print(f"  Citation embed: rewrite failed — {e}")
                continue

        except Exception as e:
            print(f"  Citation lookup: API error — {e}")
            continue

    remaining = len(claims) - resolved
    print(f"  Citation lookup complete: {resolved} resolved, "
          f"{remaining} remain UNCONFIRMED.")
    return article_text


def _normalize_bullet_cluster_bold(article_text: str) -> str:
    """
    Deterministic pass: enforce **Item Name:** bold on every bullet cluster line.

    A bullet cluster line is a list item (- or •) where:
      - The text before the first colon is 1-6 words and not yet bold-wrapped.
      - The colon is followed by a space and then sentence content (i.e. it is a
        label colon, not a URL colon or mid-sentence colon).

    Transforms:
      - Dry fog systems: Achieve...       →  - **Dry fog systems:** Achieve...
      - Ultrasonic technology: Delivers...→  - **Ultrasonic technology:** Delivers...

    Leaves already-bold lines untouched:
      - **Steam humidification:** ...     →  unchanged
    """
    import re

    fixed = 0
    lines = article_text.split("\n")
    result = []
    for line in lines:
        # Match bullet lines: optional whitespace, bullet marker, optional space, content
        m = re.match(r'^(\s*[-•]\s+)(.*)', line)
        if m:
            indent = m.group(1)
            content = m.group(2)
            # Skip if already starts with bold marker
            if not content.startswith("**"):
                # Look for a label colon: 1-6 words then ": " (not "://" for URLs)
                label_m = re.match(r'^([A-Za-z][^:\n]{2,50}):\s+(\S)', content)
                if label_m:
                    label = label_m.group(1).rstrip()
                    word_count = len(label.split())
                    # Only bold short labels (item names), not long prose fragments
                    if 1 <= word_count <= 6:
                        new_content = f"**{label}:** " + content[len(label) + 2:]
                        line = indent + new_content
                        fixed += 1
        result.append(line)

    if fixed:
        print(f"  Bullet cluster normaliser: bolded {fixed} unformatted item name(s).")
    return "\n".join(result)


def _resolve_cta_links(article_text: str, profile: dict) -> str:
    """
    Post-processing pass: find CTA anchor phrases in the article body that were
    written as plain prose (no HYPERLINK marker) and wrap them in the correct
    [HYPERLINK: anchor | url] format.

    The CTA URL is read from the client profile's cta_contact_url field, falling
    back to https://www.smartfog.com/contact-us/ for Smart Fog and
    https://veriheal.com/get-card/ for Veriheal. Only applies to sentences that
    contain a clear reader-action invitation — not every mention of Smart Fog.

    Patterns are matched case-insensitively. A phrase is only wrapped if it is
    NOT already inside a [HYPERLINK: ... ] marker.
    """
    import re

    cta_url = (
        profile.get("cta_contact_url")
        or profile.get("content_plan", {}).get("cta_contact_url")
        or ""
    )
    if not cta_url:
        # Fallback by client name
        client_name = profile.get("client_name", "").lower()
        if "smart" in client_name or "fog" in client_name:
            cta_url = "https://www.smartfog.com/contact-us/"
        else:
            return article_text  # No known CTA URL for this client — skip

    domain = profile.get("domain", "")

    # Client-specific CTA anchor patterns, ordered most-specific to least-specific.
    # Veriheal patterns cover doctor-finding and MMJ card CTAs.
    # Smart Fog patterns cover assessment, consultation, and engineer CTAs.
    if "veriheal" in domain:
        _CTA_ANCHORS = [
            r"find a (?:qualified |licensed |medical marijuana |cannabis )?doctor[^,.\n]*",
            r"connect with a (?:qualified |licensed |cannabis |medical marijuana )?doctor[^,.\n]*",
            r"speak with a (?:qualified |licensed |cannabis |medical marijuana )?(?:doctor|physician)[^,.\n]*",
            r"get (?:your )?(?:medical marijuana|MMJ|cannabis) card[^,.\n]*",
            r"apply for (?:your )?(?:medical marijuana|MMJ|cannabis) card[^,.\n]*",
            r"schedule (?:a|your) (?:free )?(?:medical marijuana |cannabis )?(?:evaluation|appointment|consultation)[^,.\n]*",
            r"book (?:a|your) (?:free )?(?:medical marijuana |cannabis )?(?:evaluation|appointment|consultation)[^,.\n]*",
        ]
    else:
        # Smart Fog and generic fallback
        _CTA_ANCHORS = [
            r"request a (?:free )?(?:system )?assessment[^,.\n]*",
            r"speak with a Smart Fog engineer[^,.\n]*",
            r"get a (?:humidification )?specification review[^,.\n]*",
            r"contact (?:a )?Smart Fog (?:engineers?|team)[^,.\n]*",
            r"request a (?:free )?(?:humidification )?consultation[^,.\n]*",
            r"request a (?:free )?quote[^,.\n]*",
            r"schedule a (?:free )?(?:system )?assessment[^,.\n]*",
        ]

    resolved = 0
    for pattern in _CTA_ANCHORS:
        for m in re.finditer(pattern, article_text, re.IGNORECASE):
            phrase = m.group(0).rstrip(" .,;:")
            start = m.start()
            # Skip if already inside a [HYPERLINK: ... ] block
            preceding = article_text[max(0, start - 12):start]
            if "[HYPERLINK:" in preceding:
                continue
            # Skip if the phrase itself contains a HYPERLINK marker already
            if "[HYPERLINK:" in phrase:
                continue
            replacement = f"[HYPERLINK: {phrase} | {cta_url}]"
            article_text = article_text[:start] + replacement + article_text[start + len(phrase):]
            resolved += 1
            print(f"  CTA resolver: linked '{phrase}' → {cta_url}")
            break  # Re-scan from scratch after each replacement to keep offsets valid

    if resolved == 0:
        print("  CTA resolver: no unlinked CTA phrases found.")
    return article_text


def _embed_missing_links(
    article_text: str,
    brief_text: str,
    client: anthropic.Anthropic,
    domain: str = "veriheal.com",
) -> str:
    """
    Second-pass revision: detects internal links specified in the brief that
    are absent from the generated article, then asks Claude to embed them
    mid-sentence within existing prose. Returns the (possibly revised) article text.
    Skips the API call entirely if all links are already present.
    """
    import re

    # 1. Extract anchor+URL pairs from the brief's internal links table rows.
    #    _read_brief renders table rows as "Anchor Text | Link To" lines.
    link_pairs: list[tuple[str, str]] = []
    in_links_section = False
    for line in brief_text.split("\n"):
        stripped = line.strip()
        if "Internal Links to Add" in stripped:
            in_links_section = True
            continue
        if in_links_section:
            if any(marker in stripped for marker in ["Disclaimer", "NeuronWriter", "---"]):
                break
            if "|" in stripped and "Anchor Text" not in stripped and "Link To" not in stripped:
                parts = [p.strip() for p in stripped.split("|")]
                if len(parts) >= 2:
                    anchor, url_val = parts[0], parts[1]
                    # Normalise relative paths to full URLs
                    if url_val.startswith("/"):
                        url_val = f"https://www.{domain}" + url_val
                    if anchor and url_val and url_val.startswith("http"):
                        link_pairs.append((anchor, url_val))

    if not link_pairs:
        return article_text

    # 2. Check which URLs are absent from the article text.
    missing: list[tuple[str, str]] = [
        (anchor, url_val)
        for anchor, url_val in link_pairs
        if url_val not in article_text
    ]

    if not missing:
        return article_text

    print(f"  Second-pass link injection: {len(missing)} missing link(s) detected.")
    for anchor, url_val in missing:
        print(f"    - {anchor} | {url_val}")

    # 3. Build the revision prompt.
    missing_block = "\n".join(
        f"  [HYPERLINK: {anchor} | {url_val}]"
        for anchor, url_val in missing
    )
    revision_prompt = f"""The following internal links are missing from this article. \
Embed each one mid-sentence within an existing sentence that discusses the same topic. \
Use the format [HYPERLINK: anchor text | url] exactly — do not use markdown link syntax.

Rules:
- Anchor text describes the destination topic in 2–5 words
- The link sits mid-sentence on the phrase most relevant to the destination
- Never start a sentence with the anchor text
- Never write a new sentence around the link — find an existing sentence and embed within it
- Rewrite the surrounding sentence slightly if needed to make the anchor text fit naturally
- INTERNAL LINK ANCHOR TEXT: Anchor text must accurately reflect the destination page topic.
  If the destination is about choosing rolling papers, the anchor must say something like
  "choosing the right rolling paper" — not "rolling papers alternatives" which implies a
  different intent. Verify: does this phrase describe what the linked page is actually about?
- DUPLICATE LINKS: Each URL may only appear once in the article. If a URL is already
  present in the article, do not embed it again — skip it.
- LINK PLACEMENT: Do not place links in the conclusion section — that section
  may only contain the CTA link. All other links must appear in body sections
  or FAQ answers.
- FAQ QUESTIONS: Never hyperlink FAQ question text — links belong in FAQ answers
  only, not on the question itself.

Published examples of correct embedding:
  "the [HYPERLINK: way you use cannabis | https://www.example.com/blog/cannabis-consumption-methods/] plays a central role in how it may affect aging"
  "Cannabis is often studied for its role in [HYPERLINK: managing pain conditions | https://www.example.com/blog/pain-relief/] linked to aging"
  "you can [HYPERLINK: find a doctor today | https://www.example.com/find-a-doctor/] to get support suited to your needs"

Return the full article text with the links embedded and nothing else.

MISSING LINKS:
{missing_block}

ARTICLE:
{article_text}"""

    # 4. Call Claude and return the revised text.
    response = client.messages.create(
        model=config.ANTHROPIC_MODEL,
        max_tokens=8000,
        messages=[{"role": "user", "content": revision_prompt}],
    )
    revised = ""
    for block in response.content:
        if hasattr(block, "text") and block.text:
            revised += block.text

    if revised.strip():
        print("  Second-pass revision complete.")
        return revised
    else:
        print("  WARNING: Second-pass returned empty response — using original article.")
        return article_text


def _embed_missing_external_links(
    article_text: str,
    client: anthropic.Anthropic,
    domain: str = "veriheal.com",
) -> str:
    """
    Pre-flight check: scans the article for HYPERLINK markers pointing to
    non-client domains. If none are found, makes a targeted Claude call
    to inject 2-3 external links into existing body paragraphs.
    Returns the (possibly revised) article text.
    """
    import re

    # Count confirmed external links (real https:// domain, not client domain, not UNCONFIRMED)
    _escaped = re.escape(domain)
    external_pattern = re.compile(r'\[HYPERLINK:[^\]]+\|\s*https?://(?!(?:www\.)?' + _escaped + r')[^\]]+\]')
    # Also count UNCONFIRMED markers — these are valid external link placeholders
    unconfirmed_pattern = re.compile(r'\[HYPERLINK:[^\]]+\|\s*UNCONFIRMED\]', re.IGNORECASE)

    confirmed = external_pattern.findall(article_text)
    unconfirmed = unconfirmed_pattern.findall(article_text)

    if confirmed or unconfirmed:
        return article_text  # external link(s) already present (confirmed or UNCONFIRMED)

    print("  External link pre-flight: no external links found — running injection pass.")

    revision_prompt = f"""This article contains no external links to scientific or medical sources. \
You must embed 2-3 external links into the existing body paragraphs. \
Do not add new sentences — find existing sentences that state a factual claim and \
embed the link mid-sentence using [HYPERLINK: anchor | url] format. \
Use UNCONFIRMED as the url if you cannot confirm the exact URL.

EXTERNAL LINK PLACEMENT AND ANCHOR TEXT:

External links must be contextually earned — only attach a link where
the surrounding sentence makes a specific factual claim that a study
or authoritative source would directly support. Never add an external
link just to meet the minimum count. A claim like "lavender contains
linalool which may reduce anxiety" earns a link. A general statement
like "herbs have been used for centuries" does not.

Anchor text must identify the specific claim being cited with enough
detail that an editor can find the correct source. Required format:
[publication year if known] + [key finding or population studied].
Examples of acceptable anchors:
- "a 2019 PMC review on terpene interactions and the entourage effect"
- "a 2021 NIH study linking daily tobacco use to stroke risk in adults"
- "research linking linalool to reduced anxiety in clinical settings"

If the year is unknown, describe the finding specifically enough that
the source is identifiable. Vague anchors like "research suggests",
"studies show", or "smoking research" are never acceptable.

Mark uncertain URLs as UNCONFIRMED but always write the anchor text
as if the editor needs to use it to find the source themselves.

Never write '[URL REQUIRED]' anywhere in the article. Use
[HYPERLINK: anchor | UNCONFIRMED] only.

EMBEDDING RULES:
- Use [HYPERLINK: anchor | url] format only — not markdown [text](url)
- Link sits mid-sentence on the factual claim, never at sentence start
- Never write a new sentence — find an existing one

Published examples:
  "A [HYPERLINK: 2022 study on epigenetic aging in cannabis users | https://pubmed.ncbi.nlm.nih.gov/36289503/] found that heavy users showed accelerated biological aging"
  "Research shows that smoking [HYPERLINK: accelerates visible skin aging | https://pubmed.ncbi.nlm.nih.gov/11966688/] by reducing collagen production"
  "Linalool may [HYPERLINK: research suggesting linalool reduces anxiety | UNCONFIRMED] by modulating GABA receptors in the brain"

Return the full article text with external links embedded and nothing else.

ARTICLE:
{article_text}"""

    response = client.messages.create(
        model=config.ANTHROPIC_MODEL,
        max_tokens=8000,
        messages=[{"role": "user", "content": revision_prompt}],
    )
    revised = ""
    for block in response.content:
        if hasattr(block, "text") and block.text:
            revised += block.text

    if revised.strip():
        injected = external_pattern.findall(revised)
        print(f"  External link injection complete: {len(injected)} external link(s) embedded.")
        return revised
    else:
        print("  WARNING: External link injection returned empty response — using original article.")
        return article_text


def _build_metadata_table(brief_text: str, url: str) -> str:
    """
    Deterministically builds the article metadata table from the already-generated
    brief and the target URL. Returns a pipe-formatted table string to prepend to
    article_text before docx rendering. Never calls Claude.

    Parses from brief_text:
      Title      — "Revised Title Tag ..." line
      Meta desc  — "Revised Meta Description ..." line
      H1         — "Revised H1 ..." line
    Derives:
      Slug       — path segment after /blog/ in the URL
    """
    import re

    def _extract_after(text: str, *prefixes: str) -> str:
        for line in text.split("\n"):
            s = line.strip()
            for prefix in prefixes:
                if s.lower().startswith(prefix.lower()):
                    value = s[len(prefix):].strip()
                    # Strip trailing character-count annotations e.g. "(58 characters)"
                    value = re.sub(r'\s*\(\d+\s+char\w*\).*$', '', value).strip()
                    if value:
                        return value
        return ""

    # Veriheal brief uses "Revised Title Tag / Revised H1" prefixes.
    # Smart Fog brief uses table rows rendered as "Meta Title | value" / "H1 | value".
    meta_title = _extract_after(
        brief_text,
        "Revised Title Tag", "Revised Title",
        "Meta Title |", "| Meta Title |",
    )
    meta_desc = _extract_after(
        brief_text,
        "Revised Meta Description", "Revised Meta Desc",
        "Meta Description |", "| Meta Description |",
    )
    h1 = _extract_after(
        brief_text,
        "Revised H1",
        "H1 |", "| H1 |",
    )

    slug_match = re.search(r'/(?:blog|insights)/([^/?#]+)/?$', url)
    slug = slug_match.group(1) if slug_match else url.rstrip("/").split("/")[-1]

    return (
        "| Field            | Value |\n"
        "|------------------|-------|\n"
        f"| Title            | {meta_title} |\n"
        f"| Meta description | {meta_desc} |\n"
        f"| Slug             | {slug} |\n"
        f"| H1               | {h1} |"
    )


def generate_article(
    url: str,
    article_data: dict,
    brief_path: str,
    keyword_data: dict,
    nlp_terms: dict,
    output_dir: str,
    profile: dict | None = None,
) -> str:
    """
    Generate the Phase 2 optimised article and save it as a .docx file.
    Returns the absolute path to the generated file.

    Args:
        profile: Client profile dict from config.load_client_profile().
                 Controls which system prompt, article rules, and context
                 files are loaded. Falls back to safe defaults if None.
    """
    _profile = profile or {}
    _domain = _profile.get("domain", "veriheal.com")
    _client_name = _profile.get("client_name", "Content")
    print("  Reading brief...")
    brief_text = _read_brief(brief_path)

    # Build metadata table deterministically from the brief — before calling Claude
    metadata_table = _build_metadata_table(brief_text, url)
    print("  Metadata table built from brief.")

    print("  Building article prompt...")
    system_prompt = _load_system_prompt(_profile)
    context = load_all_context(config, _profile)
    context_text = format_context_for_prompt(context, _profile)[:12000]
    user_message = _build_user_message(url, article_data, brief_text, keyword_data, nlp_terms, context_text, _profile)

    print("  Calling Claude API for article generation...")
    client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)
    response = client.messages.create(
        model=config.ANTHROPIC_MODEL,
        max_tokens=8000,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    )

    article_text = ""
    for block in response.content:
        if hasattr(block, "text") and block.text:
            article_text += block.text

    if not article_text.strip():
        raise RuntimeError("Claude returned an empty article — check API response.")

    print("  Article received from Claude.")

    # Normalise any markdown links Claude produced into [HYPERLINK:] format so that
    # QA, second-pass injection, and the external-link pre-flight all operate on a
    # consistent representation.  The docx renderer does this at write-time, but
    # we need it earlier so downstream checks are not confused.
    import re as _re_norm
    article_text = _re_norm.sub(
        r'\[([^\]]+)\]\((https?://[^\)]+)\)',
        lambda m: f"[HYPERLINK: {m.group(1)} | {m.group(2)}]",
        article_text,
    )

    # Strip any metadata table Claude may have written despite the instruction not to.
    # article_text stays as pure body throughout all Claude revision passes so that
    # Claude cannot accidentally drop our deterministic table on a return call.
    _meta_pipe_markers = (
        "| Meta description", "| Meta Description", "| Meta Title",
        "| Slug", "| Title", "| H1",
    )
    if any(m in article_text[:3000] for m in _meta_pipe_markers):
        _lines = article_text.split("\n")
        _skip = 0
        for _ln in _lines:
            if _ln.strip().startswith("|") or not _ln.strip():
                _skip += 1
            else:
                break
        article_text = "\n".join(_lines[_skip:]).lstrip()
        print("  Stripped Claude-generated metadata table (will be replaced by brief-derived table).")

    # Citation resolution pass 1: resolve UNCONFIRMED markers and unlinked
    # factual claims in the original Claude-generated article.
    article_text = _resolve_external_citations(article_text, client, article_data)

    # Second-pass: deterministically embed any internal links the first pass missed.
    # This pass may introduce new [HYPERLINK: anchor | UNCONFIRMED] markers.
    article_text = _embed_missing_links(article_text, brief_text, client, domain=_domain)

    # Citation resolution pass 2: resolve any UNCONFIRMED markers introduced
    # by the embedding pass above, before the external injection runs.
    article_text = _resolve_external_citations(article_text, client, article_data)

    # Third-pass: inject external links if the first two passes produced none.
    # This pass may also introduce new [HYPERLINK: anchor | UNCONFIRMED] markers.
    article_text = _embed_missing_external_links(article_text, client, domain=_domain)

    # Citation resolution pass 3: resolve any UNCONFIRMED markers introduced
    # by the external injection pass above.
    article_text = _resolve_external_citations(article_text, client, article_data)

    # Bullet cluster normaliser: enforce **Item Name:** bold on every cluster bullet.
    # Runs after all Claude passes so inconsistent bolding is fixed deterministically.
    article_text = _normalize_bullet_cluster_bold(article_text)

    # CTA resolver: wrap any plain-prose CTA phrases in [HYPERLINK: | ] format.
    # Runs after all Claude passes so it catches CTAs the model wrote without markup.
    article_text = _resolve_cta_links(article_text, _profile)

    # Strip brief annotation comments that may have leaked into the article body.
    # Patterns like (New H2 - insert after energy comparison) are meant for editors,
    # not readers. Remove them deterministically before writing the docx.
    import re as _re_annot
    article_text = _re_annot.sub(r'\(\s*New H[1-6][^)]*\)', '', article_text)
    article_text = _re_annot.sub(r'\(\s*[Ii]nsert\s+[^)]{0,80}\)', '', article_text)
    article_text = _re_annot.sub(r'\*\(\s*New H[1-6][^)]*\)\*', '', article_text)

    # Prepend the deterministic metadata table AFTER all Claude revision passes so
    # that no subsequent Claude call can accidentally drop it.
    article_text = metadata_table + "\n\n" + article_text

    # UNCONFIRMED anchor quality check — warn if anchor is vague (no year, < 6 words)
    import re as _re
    for _marker in _re.findall(r'\[HYPERLINK:([^\|]+)\|\s*UNCONFIRMED\]', article_text):
        _anchor = _marker.strip()
        _word_count = len(_anchor.split())
        _has_number = bool(_re.search(r'\d', _anchor))
        if _word_count < 6 or not _has_number:
            print(f"  WARNING anchor quality: UNCONFIRMED link has vague anchor ({_word_count} words, no year): '{_anchor}'")

    # Deterministic pre-check: count confirmed external links before calling Claude.
    # A confirmed external link is [HYPERLINK: anchor | url] where url starts with
    # "http", does not contain the client domain, and does not contain "UNCONFIRMED".
    _ext_link_pat = _re.compile(r'\[HYPERLINK:([^\|]+)\|([^\]]+)\]')
    confirmed_external_count = sum(
        1 for m in _ext_link_pat.finditer(article_text)
        if m.group(2).strip().startswith("http")
        and _domain not in m.group(2)
        and "UNCONFIRMED" not in m.group(2)
    )
    print(f"  Confirmed external links (pre-QA): {confirmed_external_count}")

    # Deterministic pre-check: detect bare-domain external links.
    # A bare-domain link is a confirmed external [HYPERLINK] whose URL path
    # is empty, "/" only, or absent — e.g. https://pubmed.ncbi.nlm.nih.gov
    # with no article ID.
    from urllib.parse import urlparse as _urlparse
    bare_domain_links: list[tuple[str, str]] = []
    for m in _ext_link_pat.finditer(article_text):
        _bd_anchor = m.group(1).strip()
        _bd_url = m.group(2).strip()
        if (
            _bd_url.startswith("http")
            and _domain not in _bd_url
            and "UNCONFIRMED" not in _bd_url
            and _urlparse(_bd_url).path.strip("/") == ""
        ):
            bare_domain_links.append((_bd_anchor, _bd_url))
    if bare_domain_links:
        print(f"  Bare-domain external links (pre-QA): {len(bare_domain_links)}")
        for _anchor, _url in bare_domain_links:
            print(f"    - {_anchor} | {_url}")

    print("  Running QA check...")
    _article_rules_for_qa = _load_article_rules(_profile)
    qa_result = _run_qa_check(
        article_text=article_text,
        brief_text=brief_text,
        nlp_terms=nlp_terms,
        keyword_data=keyword_data,
        confirmed_external_count=confirmed_external_count,
        bare_domain_links=bare_domain_links,
        client=client,
        profile=_profile,
        article_rules=_article_rules_for_qa,
    )

    if qa_result["critical_failures"]:
        # Save raw article text for inspection when QA fails
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        slug = url.rstrip("/").split("/")[-1][:40]
        raw_path = os.path.join(output_dir, f"_raw_article_{slug}_{timestamp}.txt")
        with open(raw_path, "w", encoding="utf-8") as _rf:
            _rf.write(article_text)
        print(f"\n  Raw article saved for inspection: {os.path.basename(raw_path)}")
        print("\n  QA FAILED — Article not saved. Critical failures:")
        for failure in qa_result["critical_failures"]:
            print(f"    - {failure}")
        print("\n  Full QA report:")
        print(qa_result["report"])
        print("\n  Fix these issues in the pipeline or system prompt and re-run.")
        raise RuntimeError("Article failed QA critical checks — not saved.")

    print(f"  QA verdict: {qa_result['verdict']}")

    # Save as .docx
    # Filename uses the article topic/title for readability.
    import re as _re_fn
    _raw_title = (
        article_data.get("content_plan_topic", "")
        or article_data.get("title", "")
        or article_data["url"].rstrip("/").split("/")[-1]
    )
    _safe_title = _re_fn.sub(r'[\\/*?:"<>|]', '', _raw_title).strip()[:80]
    filename = f"Article - {_safe_title}.docx" if _safe_title else f"article_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    print(f"  Article output: {filename}")
    output_path = os.path.join(output_dir, filename)

    # Normalise em dashes to hyphens before docx rendering
    article_text = article_text.replace(" — ", " - ").replace("—", "-")

    print(f"  Writing article to {output_path}...")
    _parse_and_write_doc(article_text, url, output_path, "", client_name=_client_name)
    print(f"  Article saved: {filename}")
    return output_path
