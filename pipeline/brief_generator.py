"""
brief_generator.py — Generates the Phase 1 content brief as a .docx file.
Client-agnostic: all editorial settings come from the client profile.
"""

from __future__ import annotations
import os
import json
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic
import config
from pipeline.context_loader import load_all_context, format_context_for_prompt


def _load_system_prompt(profile: dict) -> str:
    sp_path = profile.get("_system_prompt_path", "")
    if not sp_path or not os.path.exists(sp_path):
        raise FileNotFoundError(f"System prompt not found: {sp_path}")
    with open(sp_path, "r", encoding="utf-8") as f:
        return f.read()


def _load_brief_format(profile: dict) -> str:
    fmt_path = profile.get("_brief_format_path", "")
    if not fmt_path or not os.path.exists(fmt_path):
        raise FileNotFoundError(f"Brief format not found: {fmt_path}")
    with open(fmt_path, "r", encoding="utf-8") as f:
        return f.read()


def _format_keyword_data(keyword_data: dict) -> str:
    if not keyword_data:
        return "No Ahrefs keyword data available."
    lines = []

    # ahrefs_client returns primary_keyword as a plain string
    primary_kw = keyword_data.get("primary_keyword", "")
    primary_sv = keyword_data.get("primary_sv", "N/A")
    primary_kd = keyword_data.get("primary_kd", "N/A")
    if primary_kw:
        lines.append(f"Primary keyword: {primary_kw}")
        lines.append(f"  Search volume: {primary_sv}")
        lines.append(f"  Keyword difficulty: {primary_kd}")

    secondary = keyword_data.get("secondary_keywords", [])
    if secondary:
        lines.append("\nSecondary keywords:")
        for kw in secondary:
            if isinstance(kw, dict):
                lines.append(f"  - {kw.get('keyword', '')} (SV: {kw.get('sv', 'N/A')}, KD: {kw.get('kd', 'N/A')})")
            else:
                lines.append(f"  - {kw}")

    cannib = keyword_data.get("cannibalization_urls", [])
    if cannib:
        lines.append("\nCannibalization risk URLs:")
        for u in cannib:
            lines.append(f"  - {u}")

    slug = keyword_data.get("slug_assessment", {})
    if slug:
        lines.append(f"\nSlug assessment: {slug.get('recommendation', '')}")

    top_urls = keyword_data.get("top_competitor_urls", [])
    if top_urls:
        lines.append("\nTop competitor URLs:")
        for i, u in enumerate(top_urls, 1):
            lines.append(f"  {i}. {u}")

    serp = keyword_data.get("serp_overview", [])
    if serp:
        lines.append("\nSERP overview (top competitors for primary keyword):")
        for s in serp:
            lines.append(
                f"  {s.get('position')}. {s.get('url')} "
                f"(DR: {s.get('dr')}, Traffic: {s.get('traffic')})"
            )

    return "\n".join(lines) if lines else "Ahrefs returned no usable fields."


def _format_competitor_data(competitor_data: list[dict]) -> str:
    if not competitor_data:
        return "No competitor data fetched."
    lines = []
    for i, comp in enumerate(competitor_data[:5], 1):
        url = comp.get("url", "Unknown")
        word_count = comp.get("word_count", "N/A")
        title = comp.get("title", "N/A")
        headings = comp.get("headings", comp.get("h2s", []))
        h2_list = [h.get("text", h) if isinstance(h, dict) else h for h in headings if h]
        lines.append(f"\nCompetitor {i}: {url}")
        lines.append(f"  Title: {title}")
        lines.append(f"  Word count: {word_count}")
        if h2_list:
            lines.append(f"  H2s: {' / '.join(h2_list[:8])}")
        body = comp.get("body_text", "")
        if body:
            lines.append(f"  Content preview: {body[:500]}")
    return "\n".join(lines)


def _build_user_message(
    url: str,
    article_data: dict,
    gsc_data: dict | None,
    keyword_data: dict,
    competitor_data: list[dict],
    nlp_terms: dict,
    context_text: str,
    internal_link_pool: str,
    brief_format: str,
    profile: dict | None = None,
) -> str:
    _profile = profile or {}
    _client_name = _profile.get("client_name", "Client")
    _is_net_new = article_data.get("is_net_new", False)

    gsc_summary = "No GSC data provided."
    if gsc_data:
        queries = gsc_data.get("queries", [])[:10]
        top_queries = "\n".join([
            f"  - {q['query']}: {q['clicks']} clicks, {q['impressions']} impressions, CTR {q['ctr']}%, position {q['position']}"
            for q in queries
        ])
        warnings = gsc_data.get("warnings", [])
        gsc_summary = f"Top queries:\n{top_queries}"
        if warnings:
            gsc_summary += f"\nWarnings: {warnings}"

    if not nlp_terms.get("enabled"):
        neuronwriter_note = (
            "NeuronWriter integration is PENDING. "
            "Omit the NeuronWriter section from the brief entirely."
        )
    elif not nlp_terms.get("terms"):
        neuronwriter_note = "NeuronWriter returned no terms for this keyword."
    else:
        terms = nlp_terms["terms"]
        target_wc = nlp_terms.get("target_word_count", "N/A")
        terms_formatted = "\n".join([
            f"  - {t['term']}: {t['usage_range']} (used by {t['usage_pc']}% of top pages)"
            for t in terms
        ])
        paa = "\n".join([f"  - {q}" for q in nlp_terms.get("paa_questions", [])])
        neuronwriter_note = f"""NeuronWriter analysis complete.
Target word count: {target_wc}
NLP terms (insert naturally, do not stuff):
{terms_formatted}
PAA questions:
{paa}"""

    keyword_block = _format_keyword_data(keyword_data)
    competitor_block = _format_competitor_data(competitor_data)

    # Content plan context block — injected for net-new (content-plan) flows
    if _is_net_new:
        _funnel = article_data.get("content_plan_funnel", "not specified")
        _cp_type = article_data.get("content_plan_type", "not specified")
        _cluster = article_data.get("content_plan_cluster", "not specified")
        _priority = article_data.get("content_plan_priority", "not specified")
        _landing_page = article_data.get("content_plan_landing_page", "not specified")
        _strategy_notes = article_data.get("content_plan_strategy_notes", "")
        _ai_visibility_notes = article_data.get("content_plan_ai_visibility_notes", "")
        _ai_block = (
            f"\nAI Visibility Notes:\n{_ai_visibility_notes}\n"
            if _ai_visibility_notes else ""
        )
        content_plan_block = (
            f"--- CONTENT PLAN CONTEXT ---\n\n"
            f"Funnel Stage: {_funnel}\n"
            f"Content Type: {_cp_type}\n"
            f"Cluster: {_cluster}\n"
            f"Priority: {_priority}\n"
            f"Associated Landing Page: {_landing_page}\n\n"
            f"Strategy Notes:\n{_strategy_notes if _strategy_notes else '(not provided)'}\n"
            + _ai_block
            + "\n"
        )
    else:
        content_plan_block = ""

    # Slug-fidelity rules only apply to URL-based existing-article flows
    if not _is_net_new:
        slug_fidelity_block = """--- SLUG FIDELITY RULE — NON-NEGOTIABLE ---

The original URL slug is the ground truth for the article's primary
subject. Every structural decision — H1, H2s, intro framing, section
topics — must serve the slug subject specifically, not a broader
category it belongs to.

Apply this rule as follows:
- Identify the primary subject from the slug by reading it literally:
  'how-to-recover-from-edibles' = recovery from edibles specifically.
  Not cannabis overconsumption. Not getting too high generally.
  Edibles.
- The H1 must name the slug subject explicitly. If the slug says
  'edibles', the word 'edibles' must appear in the H1.
- H2s must address the slug subject specifically. General cannabis
  H2s are only acceptable if they directly serve the slug subject
  (e.g. 'How Edibles Affect Your Body Differently' is acceptable
  under 'how-to-recover-from-edibles'; 'Why Cannabis Overconsumption
  Happens' is not — it erases the edibles subject).
- The primary keyword is the SEO target. If it is broader than the
  slug subject, use the slug subject as the content anchor and treat
  the keyword as secondary framing only.
- Never broaden the subject to accommodate a higher-volume keyword.
  A lower-volume keyword that matches the slug subject exactly is
  always preferable to a higher-volume keyword that drifts it.

---"""
    else:
        slug_fidelity_block = ""

    return f"""You are writing a Phase 1 content brief for {_client_name}.

All research data has already been collected and is provided below.
Do not call any tools. Do not conduct additional research.
Do not add any sections that are not in the required output format.
Do not write a session confirmation, preamble, or commentary.
Output the brief and nothing else.

--- ARTICLE DATA ---

Target URL: {url}
Title tag: {article_data.get('title')}
Meta description: {article_data.get('meta_description')}
H1: {article_data.get('h1')}
Article slug: {article_data.get('url', '').rstrip('/').split('/')[-1]}
Word count: {article_data.get('word_count')}
Headings: {json.dumps([h['text'] for h in article_data.get('headings', [])], indent=2)}
Internal links: {json.dumps(article_data.get('internal_links', [])[:15], indent=2)}
External links: {json.dumps(article_data.get('external_links', [])[:10], indent=2)}
Body (first 4000 chars):
{article_data.get('body_text', '')[:4000]}

{content_plan_block}--- GSC DATA ---

{gsc_summary}

--- AHREFS KEYWORD DATA ---

{keyword_block}

--- COMPETITOR DATA ---

{competitor_block}

--- NEURONWRITER ---

{neuronwriter_note}

--- INTERNAL LINKING POOL ---

Only recommend URLs from this list for internal links in the brief.
Do not recommend any URL not on this list.
If no relevant URL exists in this list, omit the link rather than inventing one.
INTERNAL LINK RELEVANCE: Before selecting an internal link, verify that
the destination page is topically relevant to the sentence or section where
the link will appear. The destination page must cover the same subject the
sentence is discussing — not merely share a keyword. If the destination
page covers a different aspect of the topic (e.g. types of rolling papers
vs alternatives to rolling papers), do not include it. When in doubt,
omit the link rather than include a misleading one.

{internal_link_pool}

--- PROJECT CONTEXT ---

{context_text}

--- REQUIRED OUTPUT FORMAT ---

{slug_fidelity_block}

You must follow this format exactly and produce nothing else.
Do not output a session confirmation, preamble, or any commentary.
Do not output sections named Competitive Landscape, SERP Features,
Content Guardrails, E-E-A-T Signals, Information Gain, Keyword Strategy,
Current Article Assessment, or Recommended Changes.
Those sections do not exist in this pipeline's brief format.
Output only the sections defined in the format below, in exactly this order:

{brief_format}

Begin the brief immediately with the ## header line. Nothing before it.
"""


def _write_docx(brief_text: str, url: str, output_path: str, client_name: str = "Content") -> None:
    doc = Document()

    # Document title
    title_para = doc.add_heading(f"{client_name.upper()} CONTENT BRIEF", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"URL: {url}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    doc.add_paragraph("")

    lines = brief_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip dividers
        if stripped == "---" or stripped.startswith("|---|"):
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Markdown tables — collect all rows and render as a Word table
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not all(c in "-| " for c in row):  # skip separator rows
                    table_lines.append(row)
                i += 1
            if table_lines:
                headers = [c.strip() for c in table_lines[0].strip("|").split("|")]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for j, h in enumerate(headers):
                    hdr_cells[j].text = h
                    for run in hdr_cells[j].paragraphs[0].runs:
                        run.bold = True
                for row_line in table_lines[1:]:
                    cells = [c.strip() for c in row_line.strip("|").split("|")]
                    row_cells = table.add_row().cells
                    for j, cell_text in enumerate(cells):
                        if j < len(row_cells):
                            row_cells[j].text = cell_text
                doc.add_paragraph("")
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            _add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
            i += 1
            continue

        # Numbered lists
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)":
            _add_formatted_paragraph(doc, stripped, style="List Number")
            i += 1
            continue

        # Empty lines
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # Normal paragraph — parse inline bold
        _add_formatted_paragraph(doc, stripped)
        i += 1

    doc.save(output_path)


def _add_formatted_paragraph(doc, text: str, style: str = "Normal"):
    """Add a paragraph with inline **bold** markdown rendered as real bold."""
    para = doc.add_paragraph(style=style)
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (idx % 2 == 1)
    return para


def generate_brief(
    url: str,
    article_data: dict,
    gsc_data: dict | None,
    keyword_data: dict,
    competitor_data: list[dict],
    nlp_terms: dict,
    output_dir: str,
    profile: dict | None = None,
) -> str:
    """
    Generate the Phase 1 content brief and save as .docx.
    Returns the absolute path to the generated file.

    Args:
        profile: Client profile dict from config.load_client_profile().
                 Controls which system prompt, brief format, and context
                 files are loaded. Falls back to Veriheal defaults if None.
    """
    _profile = profile or {}

    print("  Loading context files...")
    context = load_all_context(config, _profile)
    internal_link_pool = context.get("internal_link_pool", "")
    print(f"  Internal link pool: {len(internal_link_pool)} chars, {internal_link_pool.count(chr(10))+1} URLs")
    context_text = format_context_for_prompt(context, _profile)[:12000]  # cap to ~3k tokens

    print("  Loading brief format template...")
    brief_format = _load_brief_format(_profile)

    print("  Loading system prompt...")
    system_prompt = _load_system_prompt(_profile)

    print("  Building prompt...")
    user_message = _build_user_message(
        url, article_data, gsc_data, keyword_data, competitor_data,
        nlp_terms, context_text, internal_link_pool, brief_format,
        profile=_profile,
    )

    print("  Calling Claude API...")
    client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)

    try:
        response = client.messages.create(
            model=config.ANTHROPIC_MODEL,
            max_tokens=8000,
            system=system_prompt,
            messages=[{"role": "user", "content": user_message}],
        )
    except Exception as e:
        print(f"API ERROR TYPE: {type(e)}")
        print(f"API ERROR FULL: {e}")
        raise

    # Extract text from response — may include MCP tool use blocks
    brief_text = ""
    for block in response.content:
        if hasattr(block, "text") and block.text:
            brief_text += block.text

    if not brief_text.strip():
        raise RuntimeError("Claude returned an empty brief — check API response.")

    print("  Brief received from Claude.")

    # Save as .docx
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    slug = url.rstrip("/").split("/")[-1][:40]
    filename = f"brief_{slug}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)

    print(f"  Saving brief to {output_path}...")
    _write_docx(brief_text, url, output_path, client_name=_profile.get("client_name", "Content"))
    print(f"  Brief saved: {filename}")
    return output_path
